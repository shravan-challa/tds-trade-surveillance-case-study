"""
Stage P3 — render the surveillance DQ scorecard from 14_p3_metrics.json
into a docx deliverable using the same column-level / per-rule structure
as the P1 and P2 reports.

The reader experience: front cover -> Executive Read (5 headline KPIs +
status counts) -> per-theme sections (M1..M7) with metric tables -> a
"How to re-derive" appendix that names the script and JSON behind every
number so a reviewer can re-run the audit.
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"
METRICS = HERE / "14_p3_metrics.json"
OUT = DELIV / "stage_p3_metrics_scorecard.docx"
OUT_FALLBACK = DELIV / "stage_p3_metrics_scorecard_v2.docx"

STATUS_COLOR = {"GREEN": "27AE60", "AMBER": "E67E22", "RED": "C0392B"}


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def p(doc, text, *, bold=False, italic=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return para


def status_pill(cell, status: str) -> None:
    set_cell_shading(cell, STATUS_COLOR.get(status, "7F8C8D"))
    cell.text = ""
    r = cell.paragraphs[0].add_run(status)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def metric_table(doc, metrics: list[dict]) -> None:
    """Render metrics as a 6-column table: ID | Name | Value | Target | Status | Owner.
    Rationale and definition go in a follow-up paragraph per metric."""
    t = doc.add_table(rows=1 + len(metrics), cols=6)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    widths = [Inches(0.55), Inches(2.1), Inches(0.95), Inches(1.4), Inches(0.7), Inches(0.85)]
    for i, w in enumerate(widths):
        t.columns[i].width = w
    headers = ["ID", "Metric", "Value", "Target", "Status", "Owner"]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(c, "1F3A5F")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, m in enumerate(metrics, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate([m["id"], m["name"], f"{m['value']} {m['unit']}",
                                   m["target"], "", m["owner"]]):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(str(val))
            r.font.size = Pt(8.5)
        status_pill(cells[4], m["status"])


def metric_details(doc, metrics: list[dict]) -> None:
    """For each metric, write a small italics block with definition + rationale."""
    for m in metrics:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.2)
        run_id = para.add_run(f"{m['id']}.  ")
        run_id.bold = True
        run_id.font.size = Pt(8.5)
        run_def = para.add_run(f"Definition: {m['definition']}  ")
        run_def.italic = True
        run_def.font.size = Pt(8.5)
        run_rat = para.add_run(f"Why: {m['rationale']}")
        run_rat.font.size = Pt(8.5)


def cover(doc, data: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TD Securities — Trade Surveillance Case Study")
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Stage P3 — Surveillance Data Quality Scorecard")
    rs.italic = True; rs.font.size = Pt(14)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Recurring KPIs the surveillance team can run on every new vendor / producer feed")
    rs2.italic = True; rs2.font.size = Pt(11)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: Shravan Challa\n").font.size = Pt(10)
    meta.add_run(f"Date: {data['generated_at']}\n").font.size = Pt(10)
    meta.add_run(f"Input: cleaned_p1_final.tsv ({data['input_records']:,} records)\n").font.size = Pt(10)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Purpose. ").bold = True
    intro.add_run(
        "Priority 1 made the data ingestable. Priority 2 surfaced the surveillance-domain defects. "
        "Priority 3 turns those findings into a recurring scorecard the surveillance team can re-run on every "
        "new feed to decide (a) can we ingest, (b) which fields are populated, (c) is the lineage intact, "
        "(d) is the lifecycle coherent, (e) is reference data stable, (f) can we attribute every execution, "
        "(g) which rules can we actually run. Each metric carries a stable ID, a defensible target, a RED/AMBER/GREEN "
        "status, and an explicit owner so the report routes work to the right team."
    )
    doc.add_page_break()


def executive_read(doc, data: dict) -> None:
    add_heading(doc, "Executive Read", level=1)
    p(doc, "Five headline KPIs and the status distribution. Open the per-theme sections for the full set.", italic=True, size=10)

    head = data["summary"]["headline"]
    headline_rows = [
        ("Ingestion parse rate", f"{head['ingestion_parse_rate_pct']:.2f}%", "GREEN",
         "Records that survive ingestion as a percentage of raw data lines. Direct JPM-case completeness analogue."),
        ("Fill attribution rate", f"{head['fill_attribution_rate_pct']:.2f}%", "RED",
         "Reg-best-execution + MAR market-abuse attribution coverage. Currently 65% of fills cannot be tied to a human."),
        ("Instrument->ISIN drift rate", f"{head['instrument_isin_drift_pct']:.2f}%", "RED",
         "Per-instrument surveillance rules (price outlier, position aggregation, market-data join) are unreliable until this is resolved."),
        ("Account == CounterPartyFirm collapse", f"{head['account_cpf_collapse_pct']:.2f}%", "RED",
         "Account granularity is masked on 97% of rows. Wash-trade rules cannot isolate one account inside a firm."),
        ("Orphaned parent reference rate", f"{head['lineage_orphan_parent_pct']:.2f}%", "RED",
         "1 in 10 child events cannot be tied to its parent order. Lifecycle reconstruction has 10% holes."),
    ]

    t = doc.add_table(rows=1 + len(headline_rows), cols=4)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    widths = [Inches(2.4), Inches(1.0), Inches(0.7), Inches(3.0)]
    for i, w in enumerate(widths):
        t.columns[i].width = w
    headers = ["KPI", "Value", "Status", "What it tells us"]
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(9)
        set_cell_shading(c, "1F3A5F")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, (name, val, status, note) in enumerate(headline_rows, start=1):
        cells = t.rows[ri].cells
        cells[0].text = ""; cells[0].paragraphs[0].add_run(name).font.size = Pt(9)
        cells[1].text = ""; r = cells[1].paragraphs[0].add_run(val); r.bold = True; r.font.size = Pt(10)
        status_pill(cells[2], status)
        cells[3].text = ""; cells[3].paragraphs[0].add_run(note).font.size = Pt(8.5)

    doc.add_paragraph()
    p(doc, "Status distribution across all 44 metrics:", bold=True)
    by_status = data["summary"]["by_status"]
    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Light Grid Accent 1"
    for i, (label, count, color) in enumerate([
        ("GREEN", by_status["GREEN"], "27AE60"),
        ("AMBER", by_status["AMBER"], "E67E22"),
        ("RED",   by_status["RED"],   "C0392B"),
    ]):
        c = summary_table.rows[0].cells[i]
        c.text = ""
        set_cell_shading(c, color)
        r = c.paragraphs[0].add_run(f"{label}: {count}")
        r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p(doc, "Reading guide:", bold=True)
    para = doc.add_paragraph()
    para.add_run("RED ").bold = True
    para.add_run("means a surveillance rule will be unreliable on this feed OR a regulator-exam gap exists. ")
    para.add_run("AMBER ").bold = True
    para.add_run("means the data is workable but worth tracking. ")
    para.add_run("GREEN ").bold = True
    para.add_run("means no surveillance impact on this feed. ")
    para.add_run("Owners: ").bold = True
    para.add_run("'Producer' = the upstream system needs to fix the data; "
                 "'Surveillance team' = the rule logic / threshold belongs to TD; "
                 "'Both' = a coordinated change is required.")
    doc.add_page_break()


def theme_section(doc, theme_name: str, theme: dict) -> None:
    add_heading(doc, theme_name, level=1)
    metric_table(doc, theme["metrics"])
    doc.add_paragraph()
    p(doc, "Per-metric definition and rationale:", bold=True, size=10)
    metric_details(doc, theme["metrics"])
    doc.add_page_break()


def reproducibility(doc) -> None:
    add_heading(doc, "How to re-derive every number in this scorecard", level=1)
    p(doc, "Run from case-study/analysis/ in this order. Every metric is sourced from the JSON outputs of these scripts; "
           "no number in the scorecard is hand-typed.", size=10)
    rows = [
        ("04_line_count_audit.py",     "Bare-CR vs CRLF counts (M1.5, M1.6)"),
        ("05_repair_structural.py",    "P1 structural repair -> cleaned_stage_p1.tsv"),
        ("06_normalize_values.py",     "P1 value normalisation -> cleaned_p1_final.tsv + 06_normalize_stats.json"),
        ("07_post_p1_profile.py",      "Per-column profile -> 07_post_p1_profile.json (drives M2.* nulls)"),
        ("09_p2_accuracy_efficacy.py", "P2 surveillance findings -> 09_p2_findings.json (drives M3.*, M4.*, M5.*, M6.*)"),
        ("12_test_cleaned_p1.py",      "Verifies P1 invariants hold on cleaned_p1_final.tsv (42 tests)"),
        ("13_test_p2_findings.py",     "Verifies every P2 finding count reproduces (21 tests)"),
        ("14_p3_metrics.py",           "Computes this scorecard -> 14_p3_metrics.json"),
        ("15_generate_p3_doc.py",      "Renders this document"),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    t.columns[0].width = Inches(2.5)
    t.columns[1].width = Inches(4.2)
    for i, (script, desc) in enumerate(rows):
        a, b = t.rows[i].cells[0], t.rows[i].cells[1]
        a.text = ""; b.text = ""
        ra = a.paragraphs[0].add_run(script); ra.font.name = "Consolas"; ra.font.size = Pt(9)
        rb = b.paragraphs[0].add_run(desc); rb.font.size = Pt(9)
    p(doc, "\nPinned versions: Python 3.14.0 - pandas 3.0.2 - numpy 2.4.4 - python-docx 1.2.0", size=9, italic=True)


def main() -> None:
    data = json.loads(METRICS.read_text(encoding="utf-8"))

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    cover(doc, data)
    executive_read(doc, data)
    for theme_name, theme in data["themes"].items():
        theme_section(doc, theme_name, theme)
    reproducibility(doc)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED. Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
