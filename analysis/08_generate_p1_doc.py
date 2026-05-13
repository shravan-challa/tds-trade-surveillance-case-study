"""
Stage P1 — Repair Report (Word doc).

Documents what changed between Stage 1 (observation-only) and Stage P1 (fixes):
the structural repair, the value-level normalisation rules, per-rule impact,
per-column before/after stats, and the major findings (including the surprise:
much of the apparent column-shift contamination was parser-induced, not data
corruption).
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"

STAGE1 = json.loads((HERE / "01_profile_output.json").read_text(encoding="utf-8"))
REPAIR = json.loads((HERE / "05_repair_stats.json").read_text(encoding="utf-8"))
NORM   = json.loads((HERE / "06_normalize_stats.json").read_text(encoding="utf-8"))
POSTP1 = json.loads((HERE / "07_post_p1_profile.json").read_text(encoding="utf-8"))

OUT = DELIV / "stage_p1_repair_report.docx"
OUT_FALLBACK = DELIV / "stage_p1_repair_report_v2.docx"

# Which flags indicate a real defect (vs. legitimate emptiness)
DEFECT_FLAGS = {
    "FLAG_NULL_ExchangeId", "FLAG_NULL_MessageId",
    "FLAG_MT_CASE", "FLAG_MT_UNKNOWN_ENUM",
    "FLAG_DATE_REFORMATTED", "FLAG_DATE_UNPARSABLE",
    "FLAG_TT_NOT_ISO",
    "FLAG_PK_INVALID", "FLAG_PK_NULL", "FLAG_PK_DUPLICATE",
    "FLAG_ISIN_INVALID", "FLAG_SIDE_INVALID",
    "FLAG_PRICE_UNPARSABLE", "FLAG_PRICE_NEGATIVE",
    "FLAG_VOL_UNPARSABLE", "FLAG_VOL_NONPOSITIVE",
    "FLAG_ACCOUNT_FORMAT", "FLAG_CPTY_FORMAT",
    "FLAG_TRADER_FORMAT", "FLAG_EXCH_FORMAT",
}

LEGITIMATE_EMPTY_FLAGS = {
    "FLAG_NULL_LinkMessageId", "FLAG_NULL_ParentOrderId",
    "FLAG_NULL_Trader", "FLAG_NULL_Flags",
}


# ---------- low-level helpers (duplicated from 03 to keep this script self-contained) ----------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, mono=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if mono:
        run.font.name = "Consolas"
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text, mono=False):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    return p


def add_kv_table(doc, pairs):
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    t.columns[0].width = Inches(2.0)
    t.columns[1].width = Inches(4.5)
    for i, (k, v) in enumerate(pairs):
        c0 = t.rows[i].cells[0]
        c1 = t.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(str(k))
        r0.bold = True
        r0.font.size = Pt(9)
        r1 = c1.paragraphs[0].add_run(str(v))
        r1.font.size = Pt(9)


def flag_box(doc, label, severity, text):
    color = {
        "P1-CRITICAL": "C0392B", "P1-HIGH": "E67E22",
        "P2": "2980B9", "INFO": "7F8C8D", "RESOLVED": "27AE60",
    }.get(severity, "7F8C8D")
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.1)
    t.columns[1].width = Inches(5.4)
    cell_label = t.rows[0].cells[0]
    cell_text = t.rows[0].cells[1]
    cell_label.text = ""
    cell_text.text = ""
    set_cell_shading(cell_label, color)
    set_cell_shading(cell_text, "F4F6F8")
    rl = cell_label.paragraphs[0].add_run(severity)
    rl.bold = True
    rl.font.size = Pt(8)
    rl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = cell_text.paragraphs[0]
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)


# ---------- sections ----------
def build_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TD Securities — Trade Surveillance Case Study")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Stage P1 — Repair Report")
    rs.italic = True
    rs.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: Shravan Challa\n").font.size = Pt(10)
    meta.add_run(f"Date: {date.today().isoformat()}\n").font.size = Pt(10)
    meta.add_run("Companion to: Stage 1 — Data Profile Audit Trail\n").font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    pr = p.add_run("Purpose. ")
    pr.bold = True
    p.add_run(
        "This document records the Stage P1 repair pipeline applied to the synthetic trade-surveillance dataset, "
        "the rationale for each transformation, and the resulting before/after state. Each rule is referenced by a "
        "stable ID and is reproducible from the scripts in case-study/analysis/. As in Stage 1, every claim is tied "
        "to a count or a row-level signature so the result can be defended end-to-end."
    )
    doc.add_page_break()


def build_exec_summary(doc):
    add_heading(doc, "Executive Summary", level=1)

    add_kv_table(doc, [
        ("Raw data lines (universal-newline)", f"{STAGE1['raw_data_lines']:,}"),
        ("Stage 1 strict-parsed rows", f"{STAGE1['parsed_row_count']:,}  ({100*STAGE1['parsed_row_count']/STAGE1['raw_data_lines']:.2f}%)"),
        ("Stage 1 malformed rows", f"{STAGE1['csv_malformed_count']:,}  ({STAGE1['csv_malformed_pct']:.2f}%)"),
        ("Stage P1 clean rows", f"{POSTP1['p1_rows']:,}  ({POSTP1['p1_parse_rate_vs_raw_pct']:.2f}%)"),
        ("Net rows recovered by structural repair", f"+{REPAIR['recovered_rows_vs_baseline']:,}  ({REPAIR['recovery_rate_of_malformed_pct']:.2f}% of the Stage-1 malformed)"),
        ("Stage P1 quarantine", f"{REPAIR['post_repair_quarantine_rows']:,}  (rows failing 18-field structural validation after repair)"),
    ])

    add_para(doc, "Three headline findings:", bold=True)
    add_bullet(doc, "Structural repair recovers 96,125 of the 104,410 previously-malformed rows (92.06%). The remaining ~8,285 unaccounted lines are not lost — they are correctly absorbed as multi-line cell contents of parent records and no longer count as separate rows.")
    add_bullet(doc, "The apparent 'column-shift contamination' diagnosed in Stage 1 was almost entirely a parser-induced artifact, not a data-corruption defect class. After the structural repair, zero rows trigger the cross-column-shift indicators (MessageType ∉ enum, ISIN invalid, BuyOrSell ∉ {Buy,Sell}, TransactionTime non-ISO).")
    add_bullet(doc, "Two real residual defects survive the mechanical fixes and require business judgement to resolve: 17,671 rows participate in PK duplicates, and 5,944 rows have negative Price values that cannot represent a legitimate surveillance event in equity-style trade data.")
    doc.add_page_break()


def build_pipeline_overview(doc):
    add_heading(doc, "1. Pipeline Overview", level=1)
    add_para(doc,
             "The Stage P1 pipeline is three steps. Each step is a single Python script; outputs from one feed the next; "
             "no manual editing is performed at any point.", size=10)

    add_kv_table(doc, [
        ("Step 1 — Structural repair", "05_repair_structural.py — normalise line endings, parse with skipinitialspace=True to recover quoted multi-line fields, strip per-field whitespace, output strict 18-field TSV."),
        ("Step 2 — Value normalisation", "06_normalize_values.py — apply per-column rules N1–N10 (see §3). Each rule has a stable ID, a stated rationale, and a per-rule rows-affected count."),
        ("Step 3 — Post-fix profile", "07_post_p1_profile.py — re-run the Stage-1 profile dimensions against the cleaned data so before/after deltas can be quantified."),
    ])
    doc.add_page_break()


def build_structural(doc):
    add_heading(doc, "2. Step 1 — Structural Repair", level=1)

    add_heading(doc, "2.1 Two structural defects, one combined fix", level=2)
    add_para(doc,
             "Stage 1 identified two structural defects that together accounted for the 41.65% malformation rate:", size=10)
    add_bullet(doc, "Mixed line endings: 246,522 CRLF terminators plus 4,193 bare-CR terminators. Strict-LF readers miss the bare-CR lines silently.")
    add_bullet(doc, "Pervasive ' , ' (space-comma-space) separators that defeat CSV quote recognition: under RFC 4180, a quoted field is only recognised as quoted if the opening quote is the FIRST character of the field. The leading space turns the opening quote into a literal character, so embedded line breaks inside what were meant to be quoted multi-line cells became record separators instead.")

    add_heading(doc, "2.2 Repair logic", level=2)
    add_kv_table(doc, [
        ("Step A — line-ending normalisation", "Replace every \\r\\n and bare \\r with \\n at the byte level. After this step the file has exactly one line-terminator convention."),
        ("Step B — CSV parse with skipinitialspace=True", "Use Python's csv.reader so the leading whitespace before each field is consumed before the quote-detection state machine starts. Embedded \\n inside quoted cells now resolves correctly as part of the cell value, not as a record separator."),
        ("Step C — per-field strip", "Strip leading/trailing whitespace from every parsed value, which catches the trailing space (and any residual leading whitespace that survived step B)."),
        ("Step D — field-count validation", "Each parsed record must have exactly 18 fields. Records that don't are written to a quarantine file (none, after the steps above)."),
    ])

    add_heading(doc, "2.3 Outcome", level=2)
    add_kv_table(doc, [
        ("Input file bytes", f"{REPAIR['input_file_bytes']:,}"),
        ("Lines after EOL normalisation", f"{REPAIR['lines_after_eol_normalisation']:,}"),
        ("Records parsed (post-repair)", f"{REPAIR['post_repair_records_attempted']:,}"),
        ("Records passing 18-field validation", f"{REPAIR['post_repair_clean_rows']:,}"),
        ("Records quarantined", f"{REPAIR['post_repair_quarantine_rows']:,}"),
        ("Rows recovered vs Stage-1 strict-parse", f"+{REPAIR['recovered_rows_vs_baseline']:,}"),
        ("Recovery rate of Stage-1 malformed", f"{REPAIR['recovery_rate_of_malformed_pct']:.2f}%"),
        ("Final parse-rate vs raw data lines", f"{REPAIR['final_parse_rate_vs_raw_pct']:.2f}%"),
    ])
    flag_box(doc, "F-P1-STRUCT — Ingestion completeness restored", "RESOLVED",
             f"Recovers {REPAIR['recovered_rows_vs_baseline']:,} rows ({REPAIR['recovery_rate_of_malformed_pct']:.2f}% of Stage-1 malformed). "
             f"The remaining {STAGE1['raw_data_lines'] - REPAIR['post_repair_clean_rows']:,} 'missing' lines are not lost — they have been correctly identified as embedded "
             "content within parent records' quoted multi-line cells, and no longer count as separate rows.")
    doc.add_page_break()


def rule_block(doc, rule):
    add_kv_table(doc, [
        ("Rule ID", rule["rule_id"]),
        ("Column", rule["column"]),
        ("Rationale", rule["rationale"]),
        ("Action", rule["action"]),
        ("Rows affected", f"{rule['rows_affected']:,}"),
    ])
    doc.add_paragraph()


def build_value_normalisation(doc):
    add_heading(doc, "3. Step 2 — Value-Level Normalisation", level=1)

    add_para(doc,
             "Each rule is a small named transformation with: rule ID, column, rationale (the WHY), action (the WHAT), "
             "and a per-rule rows-affected count. The pipeline is deliberately conservative — defects are converted to "
             "nulls and flagged on a per-row basis rather than guessed-at. Aggressive recovery (e.g., re-aligning shifted "
             "rows by inserting missing fields) is deferred to Stage P2 once a producer review confirms the shift signature.", size=10)

    # Group rules by category
    rules = NORM["rules"]
    add_heading(doc, "3.1 N1 — Null-token unification", level=2)
    for r in rules:
        if r["rule_id"].startswith("N1."):
            rule_block(doc, r)

    add_heading(doc, "3.2 N2 — MessageType case normalisation", level=2)
    for r in rules:
        if r["rule_id"].startswith("N2"):
            rule_block(doc, r)

    add_heading(doc, "3.3 N3 — MessageDate format canonicalisation", level=2)
    for r in rules:
        if r["rule_id"].startswith("N3"):
            rule_block(doc, r)

    add_heading(doc, "3.4 N4 — TransactionTime ISO validation", level=2)
    for r in rules:
        if r["rule_id"].startswith("N4"):
            rule_block(doc, r)

    add_heading(doc, "3.5 N5 — MessageId primary-key validation", level=2)
    for r in rules:
        if r["rule_id"].startswith("N5"):
            rule_block(doc, r)

    add_heading(doc, "3.6 N6 — ISIN format validation", level=2)
    for r in rules:
        if r["rule_id"].startswith("N6"):
            rule_block(doc, r)

    add_heading(doc, "3.7 N7 — BuyOrSell enum validation", level=2)
    for r in rules:
        if r["rule_id"].startswith("N7"):
            rule_block(doc, r)

    add_heading(doc, "3.8 N8 — Price coercion", level=2)
    for r in rules:
        if r["rule_id"].startswith("N8"):
            rule_block(doc, r)

    add_heading(doc, "3.9 N9 — TotalVolume coercion", level=2)
    for r in rules:
        if r["rule_id"].startswith("N9"):
            rule_block(doc, r)

    add_heading(doc, "3.10 N10 — Identifier format checks (flag-only)", level=2)
    add_para(doc,
             "Format checks on Account, CounterPartyFirm, Trader, ExchangeId. Flag-only — values are not nulled, "
             "because a value that doesn't match the pattern may still be informative (e.g., the firm may have legitimate "
             "format variations). The flags are surfaced for review.", size=10)
    for r in rules:
        if r["rule_id"].startswith("N10."):
            rule_block(doc, r)

    add_heading(doc, "3.11 N11 — Column-shift score (row-level)", level=2)
    for r in rules:
        if r["rule_id"] == "N11":
            rule_block(doc, r)
    flag_box(doc, "F-P1-SHIFT — Column-shift was a parser artifact", "RESOLVED",
             "Zero rows trigger ≥2 cross-column shift indicators after the structural repair. The 'column shift' pattern "
             "observed in Stage 1 (ISINs in MessageType, Buy/Sell in TransactionTime, volumes in MessageId) was a "
             "downstream artifact of the unclosed-quote misparse — fields were reading from the wrong physical line, "
             "not from a wrong logical column. Now that csv.reader correctly absorbs the embedded line breaks, every "
             "value lands in its intended column. This is a meaningful simplification of the remediation surface.")
    doc.add_page_break()


def build_residual_defects(doc):
    add_heading(doc, "4. Residual Defects (Real, Post-Repair)", level=1)

    add_para(doc,
             "After structural repair and value normalisation, four classes of defect remain. These are real DQ issues, "
             "not parser artifacts, and each requires a small judgement call rather than a mechanical fix.", size=10)

    # Find specific rule counts from NORM
    by_id = {r["rule_id"]: r for r in NORM["rules"]}

    add_heading(doc, "4.1 PK duplicates", level=2)
    r = by_id.get("N5c", {})
    flag_box(doc, "F-P1-DUP — Duplicate MessageIds", "P1-HIGH",
             f"{r.get('rows_affected', 0):,} rows participate in PK duplicates. "
             "Three possible causes — replay-from-upstream OMS, residual merge artifacts from the malformed-quote contamination, "
             "or deliberate seed-time duplication. Each has a different remediation: (a) replay = dedupe and keep latest, "
             "(b) merge artifact = rebuild from raw with a tighter quote-repair pass, (c) seed = leave for the metrics stage. "
             "Recommendation: surface a duplicate report to the data producer before any dedup is run in production.")

    add_heading(doc, "4.2 Negative prices", level=2)
    r = by_id.get("N8b", {})
    flag_box(doc, "F-P1-NEGPX — Negative Price values", "P1-HIGH",
             f"{r.get('rows_affected', 0):,} rows have a negative Price after numeric coercion. In equity / OTC-derivative "
             "surveillance, a negative price is not physically meaningful and cannot trigger spoofing, layering, marking-the-close, "
             "or wash-trade rules. Set to null in the cleaned output and flagged. Producer should confirm whether negative prices "
             "are a seeded defect (most likely) or are emitted by any legitimate upstream system (e.g., commodity futures during "
             "stress events such as April-2020 WTI).")

    add_heading(doc, "4.3 Account format outliers", level=2)
    r = by_id.get("N10.ACCOUNT", {})
    flag_box(doc, "F-P1-ACCT — Account format outliers", "P2",
             f"{r.get('rows_affected', 0):,} rows have Account values that don't match the canonical pattern "
             "([A-Z]{3,5}\\d{1,3}\\.[A-Z]+). These may be legitimate format variation across desks; surfaced for review at Stage P2.")

    add_heading(doc, "4.4 ExchangeId format outliers", level=2)
    r = by_id.get("N10.EXCH", {})
    flag_box(doc, "F-P1-EXCH — ExchangeId format outliers", "P2",
             f"{r.get('rows_affected', 0):,} rows have ExchangeId values that don't match [A-Z]{{2,6}} after the N1 step removed "
             "placeholder nulls. Likely truncations or alternate codes — surfaced for review.")

    add_heading(doc, "4.5 Empty primary key", level=2)
    r = by_id.get("N5b", {})
    flag_box(doc, "F-P1-PKNULL — Empty MessageId", "P1-HIGH",
             f"{r.get('rows_affected', 0):,} rows have an empty MessageId. PK-null records cannot be deduplicated, reconciled "
             "to upstream OMS, or referenced by downstream Cancel/Fill events. These rows are kept in the cleaned output but "
             "will be excluded from PK-dependent metrics in Stage P3.")

    doc.add_page_break()


def build_per_column_delta(doc):
    add_heading(doc, "5. Per-Column Before/After", level=1)
    add_para(doc, "Stage-1 stats are computed on the raw file (with whitespace, mixed null tokens, etc.). "
                  "Stage-P1 stats are computed on the post-cleanup output. All values are after strip in both stages "
                  "for fairness of comparison.", size=10)

    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    headers = ["Column", "Stage-1 %empty", "Stage-1 nunique", "P1 %null", "P1 nunique", "Δ nunique"]
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        set_cell_shading(c, "1F3A5F")
        for paragraph in c.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for col, info in POSTP1["columns"].items():
        row = t.add_row().cells
        row[0].text = col
        row[1].text = f"{info['stage1']['pct_empty']:.2f}%"
        row[2].text = f"{info['stage1']['nunique']:,}"
        row[3].text = f"{info['p1']['pct_null']:.2f}%"
        row[4].text = f"{info['p1']['nunique']:,}"
        delta = info["delta"]["nunique_change"]
        row[5].text = f"{delta:+,}"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_para(doc,
             "Interpretation note: drops in nunique reflect the cleaning (case normalisation, null-token unification, "
             "format reformatting). Increases would indicate new variation introduced — none observed.", size=9, italic=True)
    doc.add_page_break()


def build_row_health(doc):
    add_heading(doc, "6. Row-Level Health", level=1)

    rh = POSTP1["row_health"]
    fc = rh["flag_counts"]

    defect_total = sum(v for k, v in fc.items() if k in DEFECT_FLAGS)
    legit_total = sum(v for k, v in fc.items() if k in LEGITIMATE_EMPTY_FLAGS)

    add_para(doc,
             "Every transformation produces a per-row flag. Some flags indicate a real defect (e.g., a negative Price, "
             "a duplicate PK, a case-fragmentation fix); others indicate a legitimate empty value that simply happens "
             "to share the null-detection code path (e.g., LinkMessageId is empty for parent New Order events, by design). "
             "We separate the two before computing 'row health'.", size=10)

    add_kv_table(doc, [
        ("Total rows after P1", f"{rh['total_rows']:,}"),
        ("Rows with any defect-indicative flag", f"approx via sum-by-flag (see breakdown)"),
        ("Total defect-flag firings (sum across rows; a row can contribute to multiple)", f"{defect_total:,}"),
        ("Total legitimate-empty-flag firings (informational)", f"{legit_total:,}"),
    ])

    add_heading(doc, "6.1 Defect-indicative flag counts", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Flag"
    t.rows[0].cells[1].text = "Rows"
    for k, v in sorted(fc.items(), key=lambda kv: -kv[1]):
        if k in DEFECT_FLAGS:
            row = t.add_row().cells
            row[0].text = k
            row[1].text = f"{v:,}"
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        if c is row[0]:
                            r.font.name = "Consolas"

    add_heading(doc, "6.2 Legitimate-empty flag counts (informational only)", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "Flag"
    t.rows[0].cells[1].text = "Rows"
    for k in sorted(LEGITIMATE_EMPTY_FLAGS, key=lambda x: -fc.get(x, 0)):
        v = fc.get(k, 0)
        row = t.add_row().cells
        row[0].text = k
        row[1].text = f"{v:,}"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if c is row[0]:
                        r.font.name = "Consolas"

    add_para(doc,
             "The legitimate-empty counts are kept in the output (rather than suppressed) so a future reader can verify "
             "they correspond to the expected business rules (e.g., parent New Orders have no LinkMessageId; algo-generated "
             "messages have no Trader; non-exception messages have no Flag). Confirming these business rules is part of Stage P2.", size=9, italic=True)

    doc.add_page_break()


def build_open(doc):
    add_heading(doc, "7. Open Questions / Deferred to Stage P2", level=1)
    add_bullet(doc, "Confirm the 17,671 PK duplicates are seed-time and not residual merge artifacts. If the latter, a tighter quote-repair pass on the raw file may improve recovery further.")
    add_bullet(doc, "Confirm with producer that negative Prices are seeded defects (no legitimate equity surveillance event has a negative price).")
    add_bullet(doc, "Validate the Account format pattern — the 6,386 outliers may be legitimate variation across desks or genuine corruption.")
    add_bullet(doc, "Confirm the LinkMessageId / ParentOrderId semantic distinction. Their populated rates are near-identical; if they are effectively duplicates, the schema can be simplified.")
    add_bullet(doc, "The Flags-column's only non-empty value remains 'Unmatched to Market Data from Vendor' (~6.6K rows). Promote this to a Stage-P3 surveillance-gap metric.")
    add_bullet(doc, "The Trader null rate (~65% of rows) is large. P2 should cross-tab by MessageType to verify algo/system events legitimately carry no trader.")
    doc.add_page_break()


def build_appendix(doc):
    add_heading(doc, "Appendix A — Reproducibility", level=1)
    add_bullet(doc, "case-study/analysis/05_repair_structural.py    — structural repair", mono=True)
    add_bullet(doc, "case-study/analysis/06_normalize_values.py     — value-level normalisation", mono=True)
    add_bullet(doc, "case-study/analysis/07_post_p1_profile.py      — post-P1 profile", mono=True)
    add_bullet(doc, "case-study/analysis/08_generate_p1_doc.py      — this document", mono=True)
    add_para(doc, "Running the scripts in order regenerates every artifact (TSVs, JSONs, this docx) without manual intervention.", size=10)

    add_heading(doc, "Appendix B — Severity legend", level=2)
    add_kv_table(doc, [
        ("P1-CRITICAL", "Breaks ingestion or primary-key integrity."),
        ("P1-HIGH", "Breaks downstream joins, enum validation, or schema contracts."),
        ("P2", "Accuracy / efficacy concern — defer to Stage P2."),
        ("RESOLVED", "Stage-1 defect now resolved by the P1 pipeline."),
        ("INFO", "Observation only — informational."),
    ])


def main():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    build_cover(doc)
    build_exec_summary(doc)
    build_pipeline_overview(doc)
    build_structural(doc)
    build_value_normalisation(doc)
    build_residual_defects(doc)
    build_per_column_delta(doc)
    build_row_health(doc)
    build_open(doc)
    build_appendix(doc)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED. Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
