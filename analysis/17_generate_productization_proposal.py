"""
Standalone Productization Proposal — turns the case-study bundle into a
recurring TD surveillance capability.

Audience: someone deciding whether to fund / staff this as an ongoing product.
The exec summary already names the four framings; this document expands each
with MVP scope, success metrics, owner / consumer, sequencing, and a maturity
model that maps the M1-M7 themes to a 90-day / steady-state roadmap.

Output: deliverable/02_productization_proposal.docx
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"
METRICS = HERE / "14_p3_metrics.json"
P2_FINDINGS = HERE / "09_p2_findings.json"
OUT = DELIV / "02_productization_proposal.docx"
OUT_FALLBACK = DELIV / "02_productization_proposal_v2.docx"


# -------- docx primitives ---------------------------------------------------
def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def header_cell(cell, text: str, fill: str = "1F3A5F") -> None:
    cell.text = ""
    set_cell_shading(cell, fill)
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def p(doc, text, *, bold=False, italic=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    if size:
        run.font.size = Pt(size)
    return para


def kv_cell(cell, label: str, body: str) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    rl = para.add_run(f"{label}  ")
    rl.bold = True; rl.font.size = Pt(8.5)
    rb = para.add_run(body)
    rb.font.size = Pt(8.5)


def status_pill(cell, status: str, color: str) -> None:
    set_cell_shading(cell, color)
    cell.text = ""
    r = cell.paragraphs[0].add_run(status)
    r.bold = True; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


# -------- sections ----------------------------------------------------------
def cover(doc, metrics: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Productization Proposal")
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Turning the TDS Trade Surveillance Case Study into a recurring TD capability")
    rs.italic = True; rs.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: Shravan Challa  -  {date.today().isoformat()}").font.size = Pt(10)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run("Why this exists.  ").bold = True
    intro.add_run(
        "The case study answered the take-home brief end-to-end: the synthetic feed is now ingestable "
        f"({metrics['summary']['headline']['ingestion_parse_rate_pct']:.2f}% parse rate, zero quarantine), "
        "the surveillance-domain defects are surfaced, and 44 KPIs across 7 themes operationalise the findings "
        "as recurring measurements. But the bigger value isn't the one-shot audit -- it's the pipeline itself, "
        "which is the prototype for a recurring TD surveillance capability that pays back on every new vendor / "
        "producer feed onboarded after this one."
    )

    intro2 = doc.add_paragraph()
    intro2.add_run("What this document is.  ").bold = True
    intro2.add_run(
        "Four framings for that capability, ordered by the team that would consume them. For each: the "
        "user, the MVP scope, the success metric, the cost of doing nothing, the owner. A recommended sequence "
        "for which to build first. A maturity model mapping today's M1-M7 baseline to a 90-day target and a "
        "steady state. And a short list of decisions TD leadership would need to confirm before this becomes a programme."
    )

    intro3 = doc.add_paragraph()
    intro3.add_run("Strategic anchor.  ").bold = True
    intro3.add_run(
        "The 2024 JPMorgan / Fed-OCC trade-surveillance enforcement is the case the brief named. Post-enforcement, "
        "every Tier 1 bank's surveillance team has board-level pressure to move data quality from "
        "'incident-driven postmortem' to 'forward-looking, defensible, contractual'. This proposal is one path to that."
    )
    doc.add_page_break()


def recommendation_at_a_glance(doc, metrics: dict) -> None:
    add_heading(doc, "Recommendation at a glance", level=1)
    p(doc, "Build #1 first. #2 follows naturally on the same infrastructure. #3 unlocks any time after #1 ships. "
           "#4 is the cultural shift that lands once #1 and #2 are operational.", italic=True, size=10)

    t = doc.add_table(rows=5, cols=4)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for i, w in enumerate([Inches(0.5), Inches(2.6), Inches(1.4), Inches(2.9)]):
        t.columns[i].width = w
    for i, h in enumerate(["#", "Framing", "Sequencing", "One-line value"]):
        header_cell(t.rows[0].cells[i], h)

    rows = [
        ("1", "Vendor Feed Onboarding Service", "RECOMMENDED -- start here",
         "Cuts feed-onboarding time from 4-6 weeks to days; blocks defects pre-prod."),
        ("2", "Producer Trust Scorecard", "Quarter 2 -- on top of #1",
         "Catches DQ regression in active feeds before a surveillance alert misfires."),
        ("3", "Surveillance Audit Pack", "Unlocked once #1 ships",
         "Defensibility-as-a-deliverable for regulator inquiries and internal compliance reviews."),
        ("4", "Surveillance Data SLOs", "Year 2 -- culture shift",
         "Moves DQ from incident-driven postmortem to contractual, signed by producers."),
    ]
    for ri, (num, name, seq, value) in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        cells[0].text = ""
        rn = cells[0].paragraphs[0].add_run(num); rn.bold = True; rn.font.size = Pt(11)
        cells[1].text = ""
        nm = cells[1].paragraphs[0].add_run(name); nm.bold = True; nm.font.size = Pt(9.5)
        cells[2].text = ""
        sq = cells[2].paragraphs[0].add_run(seq); sq.font.size = Pt(8.5)
        if "RECOMMENDED" in seq:
            sq.bold = True
        cells[3].text = ""
        cells[3].paragraphs[0].add_run(value).font.size = Pt(8.5)
    doc.add_paragraph()


def proposal_block(doc, num: int, name: str, fields: dict, recommended: bool = False) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    rt = title.add_run(f"#{num}.  {name}")
    rt.bold = True; rt.font.size = Pt(13)
    rt.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    if recommended:
        rec = title.add_run("    RECOMMENDED -- BUILD FIRST")
        rec.bold = True; rec.font.size = Pt(9)
        rec.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    pitch_para = doc.add_paragraph()
    pitch_para.add_run("Pitch.  ").bold = True
    pitch_para.add_run(fields["pitch"])

    t = doc.add_table(rows=8, cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    t.columns[0].width = Inches(1.7)
    t.columns[1].width = Inches(5.7)
    pairs = [
        ("User",                 fields["user"]),
        ("Job to be done",       fields["jtbd"]),
        ("MVP scope (90 days)",  fields["mvp"]),
        ("Success metric",       fields["metric"]),
        ("Cost of doing nothing", fields["cost"]),
        ("Owner / build team",   fields["owner"]),
        ("Consumer",             fields["consumer"]),
        ("Reuses from case study", fields["reuses"]),
    ]
    for i, (k, v) in enumerate(pairs):
        cells = t.rows[i].cells
        cells[0].text = ""
        rk = cells[0].paragraphs[0].add_run(k); rk.bold = True; rk.font.size = Pt(9)
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(v).font.size = Pt(9)
    doc.add_paragraph()


def four_proposals(doc) -> None:
    add_heading(doc, "The four framings", level=1)

    proposal_block(doc, 1, "Vendor Feed Onboarding Service", {
        "pitch": "Every new vendor or producer feed runs through this pipeline before it can supply production "
                 "surveillance. The output is a RAG scorecard that greenlights, conditionally onboards, or blocks "
                 "the feed. Producers receive a remediation list with named owners and SLAs. Same artefact, "
                 "parameterized per feed.",
        "user": "Surveillance Data team -- specifically the engineer or analyst who today gets handed a new feed "
                "and has to figure out, ad-hoc, whether it's safe to wire up.",
        "jtbd": "When a new feed lands, decide quickly and defensibly whether to ingest it, and if not, hand the "
                "producer a precise list of what to fix.",
        "mvp": "(a) Pipeline parameterized to run on any CSV/feed (currently hardcoded to one). "
               "(b) Onboarding gate UI: drop a file, get a scorecard back in <1 hour. "
               "(c) Producer remediation report auto-generated from the M1-M7 RED metrics. "
               "(d) Decision log: greenlight / conditional / blocked, with named approver.",
        "metric": "Mean time-to-feed-onboarding (target: <=3 business days, vs current 4-6 weeks). "
                  "Pre-prod blocked defects per feed (count). "
                  "Surveillance-rule eligibility on day 1 of production (target: >=95% on critical themes M1, M3, M6).",
        "cost": "Each new vendor onboarding today costs 4-6 weeks of ad-hoc engineering. Defects discovered post-launch "
                "trigger surveillance gaps (JPM-class exposure) and producer-relationship friction (renegotiating "
                "after the contract is signed is much harder than blocking before).",
        "owner": "Surveillance Data team builds and operates. Estimated 1 mid-level data engineer + 0.25 PM, 90 days to MVP.",
        "consumer": "Producer-management function (hands the report back to vendors); Surveillance Operations (the "
                    "downstream consumer of greenlit feeds).",
        "reuses": "100% of the case-study pipeline (analysis/01..16). The capability IS the pipeline, productionised. "
                  "The 63 reproducibility tests become the regression suite for the gate itself.",
    }, recommended=True)

    proposal_block(doc, 2, "Producer Trust Scorecard", {
        "pitch": "Recurring re-run of the onboarding gate against active production feeds. Trends every M1-M7 KPI per "
                 "producer over time. Drift below target triggers a workflow to the producer with a named remediation "
                 "owner and a date. Same play as the candidate's S1 story (Data Governance Hub: 60% reduction in "
                 "non-compliant transfers via dashboard + outreach workflow), applied to surveillance data.",
        "user": "Producer-management lead at TD Securities -- the person whose job is keeping vendor feeds healthy "
                "in production, not just at onboarding.",
        "jtbd": "Detect DQ regressions in active feeds before a surveillance alert misfires. Hold producers "
                "accountable on a recurring cadence with evidence, not anecdote.",
        "mvp": "(a) Scheduled (weekly) re-run of #1's gate against all active feeds. "
               "(b) Per-producer trend dashboards on the M1-M7 scorecard. "
               "(c) Automated outreach workflow when any KPI flips RED (Power Automate or equivalent). "
               "(d) Producer-facing portal showing their feed's current scorecard.",
        "metric": "Mean time to producer-acknowledged remediation. "
                  "Percentage of active producers with no AMBER/RED metrics for 90+ days (target: 80%). "
                  "Number of surveillance alerts that misfire because of upstream DQ drift (target: zero).",
        "cost": "DQ regressions in production go undetected until a surveillance alert fails to fire. That's the "
                "JPM-class scenario: the system never received the data it needed to surveille, and nobody noticed "
                "until the regulator did.",
        "owner": "Surveillance Data team operates the platform; Producer Mgmt owns the producer-facing relationship "
                 "and the outreach workflow.",
        "consumer": "Producer Management; Surveillance Operations; Internal Audit.",
        "reuses": "100% of #1's infrastructure plus a scheduler and a workflow layer. The dashboard work mirrors the "
                  "candidate's existing TD Power BI experience (S1, S2 stories).",
    })

    proposal_block(doc, 3, "Surveillance Audit Pack", {
        "pitch": "On-demand bundle of (cleaned data + RAG scorecard + decision register + test pass log + "
                 "methodology journal) for any date range. Hand it to regulators on inquiry, internal audit on "
                 "review, or the board on quarterly reporting. Defensibility-as-a-deliverable, not "
                 "defensibility-as-a-narrative.",
        "user": "Compliance / Internal Audit lead handling a regulator inquiry; or the surveillance team itself "
                "when a post-incident review is requested.",
        "jtbd": "Produce a defensible, signed, point-in-time evidence pack for any historical surveillance feed, "
                "fast.",
        "mvp": "(a) Date-range parameterization on the existing pipeline. "
               "(b) Single-PDF rendering of the bundle (cover + scorecard + supporting docs). "
               "(c) Sign-off workflow: surveillance lead signs the bundle as 'this is what the data looked like on date X'. "
               "(d) Archive: signed bundles are immutable + queryable by date / feed / producer.",
        "metric": "Time from inquiry to bundle-delivered (target: <=1 business day). "
                  "Re-request rate for additional evidence after delivery (target: zero -- the first bundle is complete). "
                  "Coverage: number of historical feeds for which a bundle can be produced on demand.",
        "cost": "Each audit / inquiry today is a manual scramble. Engineers reverse-engineer what the data looked "
                "like at a past point in time; defensibility is anecdotal not artefact-based. Slow + uncertain when the "
                "regulator is on the clock.",
        "owner": "Compliance + Surveillance Data joint. Compliance owns the workflow + sign-off; Surveillance Data "
                 "owns the pipeline + archive.",
        "consumer": "External regulators (OSFI, IIROC, SEC depending on jurisdiction); internal audit; board / "
                    "risk committee.",
        "reuses": "The methodology journal IS the decision register. The test pass log already exists "
                  "(scripts 12, 13). The cleaned data is already the bundle output. Glue + sign-off + archive are net-new.",
    })

    proposal_block(doc, 4, "Surveillance Data SLOs", {
        "pitch": "Define what 'surveillance data is healthy' as 7 measurable Service-Level Objectives -- the M1-M7 "
                 "themes from the P3 scorecard, signed by producers. Pages on red. Quarterly review on amber. "
                 "Moves the relationship with producers from 'we'll let you know if your data breaks something' to "
                 "'here's what fit-for-purpose means; this is what you've signed up to deliver'.",
        "user": "Surveillance Data leadership; producer leadership; platform reliability engineering function.",
        "jtbd": "Replace incident-driven DQ with contractual DQ. Make producer commitments a measurable, monitored, "
                "renegotiable document.",
        "mvp": "(a) M1-M7 themes documented as SLO definitions with target values, measurement frequency, and "
               "breach handling. "
               "(b) SLA template for producer contracts referencing the SLOs. "
               "(c) Alerting + paging integration so an SLO breach pages the on-call producer-mgmt + surveillance team. "
               "(d) Quarterly SLO review forum chaired by surveillance data leadership.",
        "metric": "Number of producers who have signed an SLA referencing the surveillance SLOs. "
                  "SLO breach count per quarter (trending down). "
                  "Mean time to SLO breach acknowledgement (target: <=4 business hours).",
        "cost": "DQ remains incident-driven. There's no contractual obligation on producers to maintain feed quality, "
                "so producers prioritise their own roadmap over surveillance's needs. Every DQ incident is a one-off "
                "negotiation rather than enforcement of a pre-agreed standard.",
        "owner": "Surveillance Data leadership owns the SLO definitions; Platform Engineering / SRE owns the alerting "
                 "infrastructure; Producer Mgmt owns the SLA conversations with vendors.",
        "consumer": "Producer Mgmt (uses SLOs in vendor negotiations); Surveillance Ops (uses breach alerts); "
                    "Compliance (uses SLO compliance as audit evidence).",
        "reuses": "M1-M7 themes already exist with defensible targets in the P3 scorecard. The case study has done the "
                  "definitional work. What's net-new is the contractual + alerting wrapper.",
    })


def maturity_model(doc, metrics: dict) -> None:
    add_heading(doc, "Maturity model -- M1-M7 today vs Q1 vs steady state", level=1)
    p(doc, "What good looks like over time, per theme. The Q1 column is what's reasonable to expect once "
           "Framing #1 (Vendor Feed Onboarding Service) is live. Steady state is the long-run target with #2-#4 in place.",
           italic=True, size=10)

    t = doc.add_table(rows=8, cols=4)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for i, w in enumerate([Inches(1.6), Inches(2.0), Inches(2.0), Inches(2.0)]):
        t.columns[i].width = w
    for i, h in enumerate(["Theme", "Today (this synthetic feed)", "Q1 -- post Framing #1", "Steady state"]):
        header_cell(t.rows[0].cells[i], h)

    rows = [
        ("M1 Ingestion completeness",
         "96.70% parse, zero quarantine. Bare-CR risk on legacy readers.",
         "100% parse on net-new feeds (gate enforces). Bare-CR normalised at producer.",
         "All active feeds at >=99.5% parse rate. No silent-drop risk anywhere."),
        ("M2 Field completeness",
         "17 of 18 fields tracked; ~12 RED (PK, Trader, Account at risk).",
         "Producer commitments on critical fields (MessageId, TransactionTime, Trader, Account).",
         "100% on critical fields; documented optionality on the rest."),
        ("M3 Lineage integrity",
         "9.92% orphaned ParentOrderId; 5.73% orphaned LinkMessageId.",
         "<2% orphan rate on net-new feeds (gate enforces parent-first ordering).",
         "<1% across all active feeds. Zero orphan-by-omission."),
        ("M4 Lifecycle coherence",
         "Time-travel on 6.06% of children; 1,748 over-fills; 1,286 fills-after-cancel.",
         "Producer SLA for parent-first ordering; over-fill cap investigated.",
         "Clean lifecycle; B3 (fill-after-cancel) treated as surveillance signal not DQ defect."),
        ("M5 Reference data stability",
         "Instrument->ISIN drift on 88% of instruments. Account collapsed to CPF on 97% of rows.",
         "Reference-data resolution playbook; producer remediation on the column-copy bug.",
         "<5% drift; granular Account values restored on >=95% of rows."),
        ("M6 Attribution coverage",
         "35% of fills carry a Trader. Reg-best-execution + MAR attribution at risk.",
         "Producer commitment to OMS backfill. >=80% Trader on fills as interim target.",
         "100% Trader attribution on all executions; documented exception process for algo flow."),
        ("M7 Rule eligibility",
         "F2 wash-trade and G1 outliers unreliable (downstream of M5).",
         "Re-baseline F2 + G1 thresholds AFTER M5 fixes ship.",
         "All P2 surveillance rules running with calibrated thresholds and known eligible-population sizes."),
    ]
    for ri, (theme, today, q1, steady) in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate([theme, today, q1, steady]):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(val)
            r.font.size = Pt(8.5)
            if ci == 0:
                r.bold = True
    doc.add_paragraph()


def open_decisions(doc) -> None:
    add_heading(doc, "Decisions TD leadership needs to confirm", level=1)
    p(doc, "Concrete things to land on before this becomes a programme. None of these are technical -- they are "
           "scope, ownership, and funding decisions.", italic=True, size=10)

    decisions = [
        ("D1.  Which surveillance system(s) does the gate cover?",
         "Today's case study used a generic schema. In production this binds to specific vendor schemas (NICE Actimize, "
         "SMARSH, Nasdaq SMARTS, in-house). Defines the parameterization scope of the pipeline."),
        ("D2.  Who is the accountable executive for surveillance data quality at TDS?",
         "The proposal works best with a single named accountable executive who owns the M1-M7 SLOs and signs off on "
         "the producer SLAs. Without one, Framing #4 doesn't land."),
        ("D3.  Producer relationship model -- contractual SLA or operational MOU?",
         "Framing #4 (SLOs) lives or dies on whether producers sign formal SLAs (vendor contracts) or operational MOUs "
         "(internal feeds from desk-side OMSs). Different governance, different teeth."),
        ("D4.  Build vs buy for the gate UI / scheduler / archive?",
         "The pipeline is the IP; the wrapping (UI, scheduler, archive, alerting) can be built or could ride on "
         "existing TD platforms (Azure Data Factory, Databricks Workflows, ServiceNow). 90-day MVP timeline depends on this."),
        ("D5.  How does this relate to existing TD enterprise data governance?",
         "If TD already has a data-governance function with its own scorecards and SLAs, this proposal needs to "
         "integrate, not duplicate. Worth a one-meeting alignment with the central data-governance team before scoping."),
        ("D6.  Audit-pack signing authority (Framing #3)?",
         "Who signs an Audit Pack as 'this is the defensible state of the data on date X'? Likely Surveillance Data "
         "lead + Compliance lead jointly, but needs explicit decision before the workflow is built."),
    ]
    for label, body in decisions:
        para = doc.add_paragraph()
        rl = para.add_run(label); rl.bold = True; rl.font.size = Pt(10)
        para_b = doc.add_paragraph()
        para_b.paragraph_format.left_indent = Inches(0.3)
        rb = para_b.add_run(body); rb.font.size = Pt(9.5)


def closing(doc) -> None:
    add_heading(doc, "Why this proposal is defensible -- the case study IS the receipts", level=1)
    p(doc, "Three things to note about how this proposal was built:", italic=True, size=10)
    bullets = [
        "Every claim about what the pipeline can do is backed by the bundle in the same folder. The 96.70% parse rate, "
        "the 44 KPIs, the M1-M7 RAG status -- all reproducible from data/synthetic_trade_data.csv via analysis/01..17.",
        "Every recommended Framing reuses infrastructure that already exists. None of the four requires net-new data "
        "engineering past parameterization + UI / scheduler / archive wrapping. The hard part (the surveillance-domain "
        "DQ logic) is done.",
        "The producer-vs-team ownership split is the same one in the P3 scorecard and the exec summary. Same RACI, "
        "same accountability matrix, applied to ongoing operations rather than a single audit.",
    ]
    for b in bullets:
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(b)
        r.font.size = Pt(10)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.add_run("Bottom line.  ").bold = True
    closing.add_run(
        "The take-home brief asked for a defensible audit. The audit is shipped. This document is what the team "
        "should consider building NEXT, given that the audit itself revealed a recurring need that today's process "
        "doesn't address. The ask is small: 90 days, one mid-level engineer, 0.25 PM, on infrastructure that already "
        "exists. The pay-off is a permanent reduction in JPM-class exposure on every new feed TDS onboards from this "
        "point forward."
    )


def main() -> None:
    metrics = json.loads(METRICS.read_text(encoding="utf-8"))

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    cover(doc, metrics)
    recommendation_at_a_glance(doc, metrics)
    four_proposals(doc)
    doc.add_page_break()
    maturity_model(doc, metrics)
    doc.add_page_break()
    open_decisions(doc)
    doc.add_paragraph()
    closing(doc)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED. Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
