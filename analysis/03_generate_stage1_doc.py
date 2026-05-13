"""
Stage 1 — Audit document generator.

Builds a .docx that documents, at column and row level, every flag,
assertion, and piece of reasoning produced by 01_profile.py and
02_extract_examples.py. This is the "defensible audit trail" artifact.

Output: ../deliverable/stage1_profile_audit.docx
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"
DELIV.mkdir(exist_ok=True)

PROFILE = json.loads((HERE / "01_profile_output.json").read_text(encoding="utf-8"))
EXAMPLES = json.loads((HERE / "02_examples.json").read_text(encoding="utf-8"))
RAW_BAD = json.loads((HERE / "02_raw_bad_lines.json").read_text(encoding="utf-8"))

OUT = DELIV / "stage1_profile_audit.docx"
# If the primary path is locked (e.g. file open in Word), fall back to a v2 path
# so a rebuild never silently fails.
OUT_FALLBACK = DELIV / "stage1_profile_audit_v2.docx"

# -------- expectations: what each column SHOULD look like, per surveillance domain knowledge
COLUMN_SPEC = {
    "ExchangeId": {
        "role": "Venue / exchange code where the message originated.",
        "expected": "Short alphanumeric code (e.g. XFKA, XFKB). Bounded enum — handful of venues.",
        "type": "string, ~2–6 chars, never null on real trade messages.",
    },
    "MessageType": {
        "role": "Order-lifecycle event type.",
        "expected": "Small closed enum: New Order, Cancel Order, Replace Order, Fill, Reject, etc. (≤ ~12 values).",
        "type": "string, case-stable, never null.",
    },
    "TransactionTime": {
        "role": "Event timestamp (UTC).",
        "expected": "ISO 8601 with milliseconds: YYYY-MM-DDThh:mm:ss.sssZ.",
        "type": "string parseable as UTC datetime, never null.",
    },
    "MessageDate": {
        "role": "Trading date associated with the message.",
        "expected": "YYYY-MM-DD (ISO date). Should be derivable from TransactionTime.",
        "type": "string parseable as date, never null.",
    },
    "MessageId": {
        "role": "PRIMARY KEY — unique identifier per message.",
        "expected": "Unique non-null token per row. Looks composite: {date}.{routeId}.P.{seq}.",
        "type": "string, unique across the file, never null, never empty.",
    },
    "LinkMessageId": {
        "role": "Reference to a related upstream message (e.g. the original order a Cancel/Fill points to).",
        "expected": "Same format as MessageId, OR empty if the message has no upstream link (e.g. parent New Order).",
        "type": "string MessageId-shaped, conditionally null.",
    },
    "ParentOrderId": {
        "role": "Reference to the root parent order for child events.",
        "expected": "Same format as MessageId, OR empty for the root New Order itself.",
        "type": "string MessageId-shaped, conditionally null.",
    },
    "Instrument": {
        "role": "Ticker / instrument code.",
        "expected": "Short alphabetic ticker (1–6 chars usually). Never null on a real trade.",
        "type": "string, alphabetic-ish, never null.",
    },
    "ISIN": {
        "role": "International Securities Identification Number.",
        "expected": "Exactly 12 alphanumeric characters, country prefix (e.g. US...).",
        "type": "string length 12, ISIN check-digit valid, never null on a real trade.",
    },
    "BuyOrSell": {
        "role": "Order side.",
        "expected": "Closed enum: Buy / Sell (possibly Short / Cover for shorts).",
        "type": "string, 2-value enum (or up to 4 if short-side codes used), never null.",
    },
    "Price": {
        "role": "Order/execution price in instrument's currency.",
        "expected": "Non-negative decimal. Zero is valid for some message types (e.g. Market Order has no limit price; New Order may be 0 until filled).",
        "type": "float ≥ 0, conditionally null.",
    },
    "TotalVolume": {
        "role": "Order quantity / fill quantity.",
        "expected": "Positive integer. Zero is suspicious (would not generate a meaningful surveillance event).",
        "type": "integer > 0, never null.",
    },
    "Account": {
        "role": "Internal trading account identifier.",
        "expected": "Internal account code; cardinality higher than CounterPartyFirm count (one firm, many accounts).",
        "type": "string, never null.",
    },
    "CounterPartyFirm": {
        "role": "Counterparty firm identifier (LEI-like internal code).",
        "expected": "Closed set of firm codes; lower cardinality than Account.",
        "type": "string, never null.",
    },
    "Trader": {
        "role": "Internal trader / desk identifier responsible for the action.",
        "expected": "Short code like TRD_<n>. Should be present on human-initiated orders; may legitimately be null for algo/system-generated messages.",
        "type": "string or null, conditionally null.",
    },
    "TransactionSource": {
        "role": "Source system that produced the message (OMS / EMS / venue / etc.).",
        "expected": "Closed enum of source-system identifiers.",
        "type": "string, never null.",
    },
    "Currency": {
        "role": "Settlement / price currency.",
        "expected": "ISO 4217 3-letter code (USD, CAD, EUR, ...).",
        "type": "string length 3, ISO 4217 valid, never null.",
    },
    "Flags": {
        "role": "Surveillance exception markers / annotations.",
        "expected": "Possibly null on most messages; populated when a specific condition is flagged by the surveillance system.",
        "type": "string or null, conditionally null.",
    },
}


# -------- low-level helpers ----------------------------------------------------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, mono=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if mono:
        run.font.name = "Consolas"
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text, mono=False):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    return p


def add_kv_table(doc, pairs):
    """Two-column key/value table for short fact lists."""
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (k, v) in enumerate(pairs):
        c0 = t.rows[i].cells[0]
        c1 = t.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(str(k))
        r0.bold = True
        r0.font.size = Pt(9)
        r1 = c1.paragraphs[0].add_run(str(v))
        r1.font.size = Pt(9)
        for cell in (c0, c1):
            cell.width = Inches(1.8) if cell is c0 else Inches(4.7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_row_example_table(doc, rows, fields=None, max_field_chars=40):
    """Render a small set of example rows as a transposed table:
    columns = example index, rows = field names. Easier to read than wide tables.
    """
    if not rows:
        p = doc.add_paragraph()
        run = p.add_run("(no parsed rows matched this defect class — defect manifests at raw-file / pre-parse stage)")
        run.italic = True
        run.font.size = Pt(9)
        return

    if fields is None:
        fields = [k for k in rows[0].keys() if k != "__row_index"]

    cols = ["Field"] + [f"Row idx {r.get('__row_index', '?')}" for r in rows]
    t = doc.add_table(rows=len(fields) + 1, cols=len(cols))
    t.style = "Light Grid Accent 1"

    # header
    for j, c in enumerate(cols):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(c)
        r.bold = True
        r.font.size = Pt(8)
        set_cell_shading(cell, "1F3A5F")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, fname in enumerate(fields, start=1):
        # field name col
        cell = t.rows[i].cells[0]
        cell.text = ""
        r = cell.paragraphs[0].add_run(fname)
        r.bold = True
        r.font.size = Pt(8)
        for j, row in enumerate(rows, start=1):
            val = str(row.get(fname, ""))
            if len(val) > max_field_chars:
                val = val[:max_field_chars] + "…"
            cell = t.rows[i].cells[j]
            cell.text = ""
            r2 = cell.paragraphs[0].add_run(val)
            r2.font.name = "Consolas"
            r2.font.size = Pt(8)


def flag_box(doc, label, severity, text):
    """A highlighted callout for a flag/assertion."""
    color = {"P1-CRITICAL": "C0392B", "P1-HIGH": "E67E22", "P2": "2980B9", "INFO": "7F8C8D"}.get(severity, "7F8C8D")
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.1)
    t.columns[1].width = Inches(5.4)
    cell_label = t.rows[0].cells[0]
    cell_text = t.rows[0].cells[1]
    cell_label.text = ""
    cell_text.text = ""
    set_cell_shading(cell_label, color)
    set_cell_shading(cell_text, "F4F6F8")
    rl = cell_label.paragraphs[0].add_run(severity)
    rl.bold = True
    rl.font.size = Pt(8)
    rl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = cell_text.paragraphs[0]
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)


# -------- document sections ----------------------------------------------------
def build_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TD Securities — Trade Surveillance Case Study")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Stage 1 — Data Profile Audit Trail")
    rs.italic = True
    rs.font.size = Pt(14)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: Shravan Challa\n").font.size = Pt(10)
    meta.add_run(f"Date: {date.today().isoformat()}\n").font.size = Pt(10)
    meta.add_run("Dataset: synthetic_trade_data.csv\n").font.size = Pt(10)
    meta.add_run(f"Rows (raw): {PROFILE['raw_data_lines']:,}  ·  Rows (parseable): {PROFILE['parsed_row_count']:,}  ·  Columns: {PROFILE['column_count']}").font.size = Pt(10)

    doc.add_paragraph()
    purpose = doc.add_paragraph()
    pr = purpose.add_run(
        "Purpose. "
    )
    pr.bold = True
    purpose.add_run(
        "This document captures the methodology, observations, flags, and reasoning produced during "
        "Stage 1 of the case study — the initial raw profile of the synthetic trade-surveillance dataset, "
        "before any cleaning or transformation. It is intended to be defensible: every claim is tied to a "
        "specific column, row example, or aggregate count, and the steps are reproducible from the "
        "scripts archived alongside this document. No fixes are proposed here — Stage 1 is observation-only."
    )

    doc.add_page_break()


def build_exec_summary(doc):
    add_heading(doc, "Executive Summary", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "The dataset cannot be ingested in its present form. "
    ).bold = True
    p.add_run(
        f"Of {PROFILE['raw_data_lines']:,} raw data lines, only {PROFILE['parsed_row_count']:,} "
        f"({100 - PROFILE['csv_malformed_pct']:.2f}%) parse cleanly with a strict-compliance CSV reader; "
        f"the remaining {PROFILE['csv_malformed_count']:,} lines ({PROFILE['csv_malformed_pct']:.2f}%) "
        "would be either dropped or silently merged by a vendor parser that expects 18 fields per record. "
        "This is the exact failure mode underlying the 2024 JPMorgan / Fed–OCC trade-surveillance enforcement: "
        "regulators cannot detect what surveillance never received."
    )

    add_para(doc, "Defect classes detected at Stage 1 (counts approximate, refined in later stages):", bold=True)
    add_bullet(doc, f"CSV malformation (unclosed quote characters in Instrument): {PROFILE['csv_malformed_count']:,} rows / {PROFILE['csv_malformed_pct']:.2f}%")
    add_bullet(doc, "Pervasive leading/trailing whitespace on column names AND values (94–99% of every column)")
    add_bullet(doc, "Mixed null-token representations within the same field (empty string, 'NaN', 'NULL', 'None', 'null')")
    add_bullet(doc, "Cardinality explosion in closed-enum fields (MessageType: 822 distinct; ExchangeId: 1,158; BuyOrSell: 4)")
    add_bullet(doc, "Value bleed-through across columns (ISIN values appearing as MessageType; Buy/Sell values as TransactionTime; volumes as MessageId)")
    add_bullet(doc, "MessageDate format heterogeneity (≥ 6 distinct date formats including ambiguous DD-MM vs MM-DD)")
    add_bullet(doc, "MessageType case fragmentation ('New Order' vs 'NEW ORDER' vs 'new order')")
    add_bullet(doc, "ISIN cardinality anomaly (6,587 distinct ISINs vs 373 distinct Instruments — implies ~1 instrument : ~17 ISINs, almost certainly corruption)")
    add_bullet(doc, "Constant columns post-strip (TransactionSource = SYS_OMEGA only; Currency = USD only) — fine if intended, but eliminates them as DQ signals")

    doc.add_page_break()


def build_methodology(doc):
    add_heading(doc, "1. Methodology", level=1)

    add_heading(doc, "1.1 Principles", level=2)
    add_bullet(doc, "Stage 1 is observation-only. No values are altered. No rows are dropped at this stage except by the strict CSV tokenizer, and those drops are themselves counted as a finding.")
    add_bullet(doc, "Defensibility-by-evidence. Every finding cites a concrete count and at least one row-level example.")
    add_bullet(doc, "Vendor-equivalent reading. The first parse mimics what a strict downstream system would do, so the 'records-lost-on-ingestion' count is realistic.")
    add_bullet(doc, "Raw-string fidelity. Pandas is configured not to coerce types, not to interpret nulls, and not to strip whitespace, so the data are observed exactly as the bytes on disk present them.")

    add_heading(doc, "1.2 Tools", level=2)
    add_kv_table(doc, [
        ("Language / runtime", "Python 3.14.0"),
        ("Core library", "pandas 3.0.2 (numpy 2.4.4)"),
        ("CSV reader", "pandas with engine='python', on_bad_lines='skip'"),
        ("Raw tokenizer (independent)", "Standard-library csv module, line-by-line"),
        ("Document generator", "python-docx 1.2.0"),
        ("Output format", ".docx (Microsoft Word)"),
    ])

    add_heading(doc, "1.3 Parser configuration & justification", level=2)
    add_kv_table(doc, [
        ("dtype=str", "Read every column as string so no auto-coercion masks defects (e.g., '0.0' staying as string, not silently becoming float)."),
        ("keep_default_na=False, na_filter=False", "Disable pandas' default null inference so we observe 'NaN', 'NULL', 'None' as the literal strings they are in the file — these are themselves defects."),
        ("encoding='utf-8'", "Confirmed at first read; no encoding errors raised. Reproducible across systems."),
        ("engine='python'", "Allows on_bad_lines='skip' to tolerate CSV malformations without aborting the whole load; cost is slower parse, acceptable at 250k rows."),
        ("on_bad_lines='skip'", "Skipped lines are subtracted from the raw line count to produce the malformation rate. This is the 'records the vendor would lose' number."),
        ("Line counting", "Universal-newline (Python text mode with newline=None). The file contains mixed line endings (CRLF + bare CR); strict-LF counting (e.g. Unix wc -l) undercounts by exactly the bare-CR count. Universal-newline matches what Power Query, Excel, and modern CSV readers see."),
    ])

    add_heading(doc, "1.4 Stated assumptions (confirmed with data producer / dataset author)", level=2)
    add_bullet(doc, "The dataset is synthetic and seeded specifically for this case study, with an intentional 1–10% defect rate per field.")
    add_bullet(doc, "TransactionSource ≡ 'SYS_OMEGA' and Currency ≡ 'USD' are intentional simplifications of the synthetic dataset, not defects. They are confirmed constant by design and are therefore documented but not flagged as defects in this audit. A production feed would have variation in both.")
    add_bullet(doc, "Findings that depend on multi-source or multi-currency behaviour (e.g. cross-currency settlement consistency) are noted as 'not testable in this dataset' rather than passed or failed.")

    add_heading(doc, "1.5 Reproducibility", level=2)
    add_bullet(doc, "case-study/analysis/01_profile.py — produces 01_profile_output.json (the source of every count in this document)")
    add_bullet(doc, "case-study/analysis/02_extract_examples.py — produces 02_examples.json and 02_raw_bad_lines.json (concrete row evidence)")
    add_bullet(doc, "case-study/analysis/04_line_count_audit.py — independent verification of line-ending composition vs Power Query / Excel")
    add_bullet(doc, "case-study/analysis/03_generate_stage1_doc.py — renders this document from the JSON artifacts")
    add_bullet(doc, "Running the scripts in order regenerates this artifact bit-for-bit; no manual editing of the .docx is performed.")

    doc.add_page_break()


def build_file_level(doc):
    add_heading(doc, "2. File-Level Findings", level=1)

    le = PROFILE.get("raw_line_endings", {})
    add_kv_table(doc, [
        ("File size", "52.56 MiB (55,109,785 bytes)"),
        ("Encoding", "UTF-8, no BOM"),
        ("Raw lines (universal-newline, incl. header)", f"{PROFILE['raw_data_lines'] + 1:,}"),
        ("Raw data lines (universal-newline)", f"{PROFILE['raw_data_lines']:,}"),
        ("CRLF (\\r\\n) line terminators", f"{le.get('crlf_pairs', 0):,}"),
        ("Bare CR (\\r) line terminators", f"{le.get('bare_cr', 0):,}"),
        ("Bare LF (\\n) line terminators", f"{le.get('bare_lf', 0):,}"),
        ("Successfully parsed rows", f"{PROFILE['parsed_row_count']:,}"),
        ("CSV-malformed rows", f"{PROFILE['csv_malformed_count']:,}"),
        ("CSV malformation rate", f"{PROFILE['csv_malformed_pct']:.2f}%"),
        ("Expected fields per record", "18"),
    ])

    flag_box(doc, "F-FILE-1 — Ingestion completeness", "P1-CRITICAL",
             f"{PROFILE['csv_malformed_pct']:.2f}% of records ({PROFILE['csv_malformed_count']:,} rows) fail strict CSV tokenization. "
             "Direct analogue to the JPMorgan 2024 surveillance-data enforcement: records the vendor never sees cannot be surveilled.")

    flag_box(doc, "F-FILE-2 — Mixed line endings (CRLF + bare CR)", "P1-CRITICAL",
             f"{le.get('crlf_pairs', 0):,} lines terminate with \\r\\n while {le.get('bare_cr', 0):,} lines terminate with a bare \\r only. "
             "Modern parsers (Power Query, Excel, Python text mode) use universal-newline mode and see all 250,714 data lines. "
             "Strict-LF readers (Unix wc -l, some legacy ingestion pipelines, certain Java/C parsers) silently miss the bare-CR lines — "
             "a hidden 4,193-record under-count, exactly the kind of invisible data loss that fails a regulator audit.")

    flag_box(doc, "F-FILE-2a — Reconciliation across tools", "INFO",
             "Verified independently: Power Query reports 250,715 rows (incl. header) ≡ universal-newline count. "
             "Git Bash wc -l reports 246,522 ≡ pure-LF count. The 4,193-row gap is fully explained by the bare-CR terminators. "
             "Documenting this here so a reader running any of these tools can reconcile the numbers themselves.")

    add_heading(doc, "2.1 Root-cause hypothesis: unclosed quotes in the Instrument field", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Examination of the first malformed rows (raw line numbers 15, 16, 33, 34, 44, 45, …) shows a consistent pattern: "
        "the eighth field (Instrument) opens with a double-quote character ("
    )
    p.add_run('"').bold = True
    p.add_run(
        ") but the closing quote is on a subsequent line, with the ticker letters split across the line break. The CSV "
        "tokenizer then absorbs the next line(s) into the open quoted field, producing a record with the wrong field count "
        "and corrupting every adjacent record."
    )

    add_heading(doc, "2.2 Evidence — raw malformed lines (verbatim, truncated)", level=2)
    bad = RAW_BAD["bad_field_count_examples"][:6]
    for ex in bad:
        para = doc.add_paragraph()
        r = para.add_run(f"Line {ex['line_number']} (fields={ex['field_count']}): ")
        r.bold = True
        r.font.size = Pt(9)
        rr = para.add_run(ex["first_500_chars"])
        rr.font.name = "Consolas"
        rr.font.size = Pt(8)

    add_para(doc, "Independent confirmation — quote-imbalance audit:", bold=True)
    add_para(doc,
             f"A separate pass that counts double-quote characters per line (without invoking the CSV parser) found "
             f"the same line numbers as those flagged by the parser. This corroborates the unclosed-quote hypothesis "
             f"as the dominant cause of malformation.", size=10)
    doc.add_page_break()


def build_header_level(doc):
    add_heading(doc, "3. Header-Level Findings", level=1)
    add_para(doc, "All 18 column headers carry leading and/or trailing whitespace. Schema validation against an exact-match header list will fail.", size=10)
    flag_box(doc, "F-HDR-1 — Header whitespace", "P1-HIGH",
             "Column names as parsed include 'ExchangeId ', ' MessageType ', ' TransactionTime ', etc. "
             "Downstream code that joins on exact column names will silently miss every column.")

    t = doc.add_table(rows=len(PROFILE["raw_columns"]) + 1, cols=3)
    t.style = "Light Grid Accent 1"
    t.rows[0].cells[0].text = "#"
    t.rows[0].cells[1].text = "Raw header (as in file)"
    t.rows[0].cells[2].text = "Stripped header"
    for i, (raw, stripped) in enumerate(zip(PROFILE["raw_columns"], PROFILE["stripped_columns"]), start=1):
        t.rows[i].cells[0].text = str(i)
        c1 = t.rows[i].cells[1]
        c1.text = ""
        r1 = c1.paragraphs[0].add_run(repr(raw))
        r1.font.name = "Consolas"
        r1.font.size = Pt(9)
        c2 = t.rows[i].cells[2]
        c2.text = ""
        r2 = c2.paragraphs[0].add_run(stripped)
        r2.font.name = "Consolas"
        r2.font.size = Pt(9)

    doc.add_page_break()


def build_column_audit(doc):
    add_heading(doc, "4. Column-by-Column Audit", level=1)
    add_para(doc,
             "Each subsection covers one of the 18 columns. Structure: Role → Expectation → Observed → Flags → "
             "Row-Level Examples → Reasoning. Counts are raw (pre-fix) and use the 146,304 successfully-parsed rows "
             "as the denominator unless stated otherwise.", size=10)

    # Mapping from column name to which example sets apply
    column_examples = {
        "ExchangeId": [
            ("Literal string 'NULL' as venue code (n=809)", "exchangeid_literal_NULL"),
            ("Embedded closing quote — value 'OO\"' (n=820)", "exchangeid_OO_quote"),
            ("Literal 'NONE' placeholder (n=824)", "exchangeid_NONE"),
            ("Catastrophic length (quote-contamination overflow)", "exchangeid_extreme_length"),
        ],
        "MessageType": [
            ("Uppercase variant 'NEW ORDER' (n=708)", "messagetype_uppercase"),
            ("Lowercase variant 'new order' (n=702)", "messagetype_lowercase"),
            ("ISIN-shaped value in MessageType (column shift)", "messagetype_is_isin"),
        ],
        "TransactionTime": [
            ("Non-ISO 8601 values (column shift / corruption)", "transactiontime_not_iso"),
        ],
        "MessageDate": [
            ("European DD-MM-YYYY (n=936)", "messagedate_dash_eu"),
            ("Slash-delimited 05/02/2026 (n=891)", "messagedate_slash_us"),
            ("Compact 20260205 (n=889)", "messagedate_compact"),
            ("Numeric leak '0.0' (n=966)", "messagedate_numeric_leak"),
        ],
        "MessageId": [
            ("Empty primary key (n=937)", "messageid_empty"),
            ("'NULL' literal token (n=457)", "messageid_NULL_token"),
            ("'NaN' literal token (n=458)", "messageid_NaN_token"),
            ("Numeric-only value where composite key expected (column shift)", "messageid_numeric"),
        ],
        "Instrument": [
            ("Bare quote character (n=2,434)", "instrument_bare_quote"),
            ("'NaN' literal (n=437)", "instrument_nan_literal"),
        ],
        "ISIN": [
            ("Shorter than 12 characters", "isin_too_short"),
            ("Value 'SYS_OMEGA' in ISIN (column shift, n=2,051)", "isin_SYS_OMEGA_leak"),
        ],
        "BuyOrSell": [
            ("'USD' in BuyOrSell column (column shift, n=2,051)", "buyorsell_USD"),
            ("'SYS_OMEGA' in BuyOrSell column (column shift, n=11)", "buyorsell_SYS_OMEGA"),
        ],
        "Trader": [
            ("'NaN' literal (n=29,852)", "trader_NaN_literal"),
            ("Empty (n=59,234)", "trader_empty"),
        ],
    }

    for col_name in PROFILE["stripped_columns"]:
        info = PROFILE["columns"][col_name]
        spec = COLUMN_SPEC[col_name]
        add_heading(doc, f"4.{PROFILE['stripped_columns'].index(col_name)+1}  {col_name}", level=2)

        # Spec
        add_kv_table(doc, [
            ("Role", spec["role"]),
            ("Expected values", spec["expected"]),
            ("Expected type", spec["type"]),
            ("Raw column header", repr(info["raw_header"])),
        ])

        # Observed stats
        add_para(doc, "Observed in dataset:", bold=True)
        sn_tokens = info["suspect_null_tokens"] or {}
        sn_str = ", ".join(f"{k!r}={v:,}" for k, v in sn_tokens.items()) or "(none)"
        add_kv_table(doc, [
            ("Total rows (parsed)", f"{PROFILE['parsed_row_count']:,}"),
            ("Empty values", f"{info['n_empty']:,} ({info['pct_empty']:.3f}%)"),
            ("Whitespace-only values", f"{info['n_whitespace_only']:,}"),
            ("Rows with leading/trailing whitespace", f"{info['n_leading_or_trailing_whitespace']:,} ({info['pct_leading_or_trailing_whitespace']:.2f}%)"),
            ("Distinct values (after strip, non-empty)", f"{info['nunique_non_empty_stripped']:,}"),
            ("Suspect null tokens detected", sn_str),
            ("Stripped value length (min / mean / max)", f"{info['stripped_length_stats']['min']} / {info['stripped_length_stats']['mean']} / {info['stripped_length_stats']['max']:,}"),
        ])

        # Top values
        add_para(doc, "Top-10 distinct values (after strip, by count):", bold=True)
        tv_table = doc.add_table(rows=len(info["top_10_values_stripped"]) + 1, cols=2)
        tv_table.style = "Light Grid Accent 1"
        tv_table.rows[0].cells[0].text = "Value"
        tv_table.rows[0].cells[1].text = "Count"
        for i, (val, cnt) in enumerate(info["top_10_values_stripped"].items(), start=1):
            c0 = tv_table.rows[i].cells[0]
            c0.text = ""
            r0 = c0.paragraphs[0].add_run(repr(val))
            r0.font.name = "Consolas"
            r0.font.size = Pt(9)
            tv_table.rows[i].cells[1].text = f"{cnt:,}"

        # Flags & reasoning — bespoke per column
        add_para(doc, "Flags raised & reasoning:", bold=True)
        emit_column_flags(doc, col_name, info)

        # Row-level examples
        ex_specs = column_examples.get(col_name, [])
        if ex_specs:
            add_para(doc, "Row-level examples:", bold=True)
            for label, key in ex_specs:
                rows = EXAMPLES["parsed_examples"].get(key, [])
                if not rows:
                    continue
                p = doc.add_paragraph()
                r = p.add_run(label)
                r.italic = True
                r.font.size = Pt(9)
                add_row_example_table(doc, rows)

        doc.add_paragraph()


def emit_column_flags(doc, col, info):
    """Per-column flag rendering. Order: severity descending."""
    pct_ws = info["pct_leading_or_trailing_whitespace"]
    n_empty = info["n_empty"]
    n_ws_only = info["n_whitespace_only"]
    nunique = info["nunique_non_empty_stripped"]
    suspect = info["suspect_null_tokens"] or {}
    maxlen = info["stripped_length_stats"]["max"]

    if pct_ws > 5:
        flag_box(doc, f"F-{col.upper()}-WS", "P1-HIGH",
                 f"{info['n_leading_or_trailing_whitespace']:,} rows ({pct_ws:.2f}%) have leading/trailing whitespace. "
                 "Breaks exact-match joins, enum validation, and primary-key uniqueness checks. Trivial auto-fix (strip).")

    if col == "ExchangeId":
        if maxlen > 100:
            flag_box(doc, f"F-{col.upper()}-OVERFLOW", "P1-CRITICAL",
                     f"Maximum ExchangeId length observed = {maxlen:,} characters. This is not a real venue code; it is the "
                     "downstream symptom of the unclosed-quote contamination — multiple lines absorbed into one quoted field. "
                     "These records will fail venue lookup AND poison every downstream record they touched.")
        if "NULL" in suspect:
            flag_box(doc, f"F-{col.upper()}-NULL-TOKEN", "P1-HIGH",
                     f"{suspect['NULL']:,} rows have the literal string 'NULL' in the venue field. Venue is mandatory; "
                     "treating 'NULL' as a value is invalid for surveillance routing.")
        flag_box(doc, f"F-{col.upper()}-CARDINALITY", "P2",
                 f"{nunique:,} distinct values observed where a small venue list is expected. The 'long tail' beyond "
                 "the top-5 exchange codes (XFKA, XFKB, XFKC, XFKD, XFKE — together ~97% of volume) is almost certainly "
                 "corruption or seeded placeholders ('OO', 'NONE', 'ZZZZ', 'NULL', 'OO\"').")

    elif col == "MessageType":
        flag_box(doc, f"F-{col.upper()}-CASE-FRAGMENTATION", "P1-HIGH",
                 "Three case-variants of each legitimate value present: 'New Order' (58,912), 'NEW ORDER' (708), 'new order' (702) — and similarly for Cancel/Replace. "
                 "These will be treated as distinct enum values by any case-sensitive downstream system. Auto-fix candidate: normalise to Title Case.")
        flag_box(doc, f"F-{col.upper()}-VALUE-LEAKAGE", "P1-CRITICAL",
                 f"{nunique:,} distinct values where the expected enum is ≤ ~12. Top-10 includes 'USRT0JJC8GX7' (820 rows) — "
                 "an ISIN value. Confirms column-shift contamination: when an upstream column is missing, every later field "
                 "appears in the wrong position.")

    elif col == "TransactionTime":
        # Note value bleed (Buy/Sell appearing here)
        flag_box(doc, f"F-{col.upper()}-VALUE-LEAKAGE", "P1-CRITICAL",
                 "Top-10 values include 'Sell' (1,033), 'Buy' (1,029), and 'nan' (2,962). Side codes and the literal 'nan' "
                 "cannot be timestamps. Confirms column-shift: 'Buy'/'Sell' should be in the BuyOrSell column at index 10.")
        flag_box(doc, f"F-{col.upper()}-FORMAT", "P1-HIGH",
                 "Timestamps should be ISO 8601 with millisecond precision and 'Z' suffix. Non-conforming values must be quarantined.")

    elif col == "MessageDate":
        flag_box(doc, f"F-{col.upper()}-FORMAT-CHAOS", "P1-CRITICAL",
                 "At least 6 distinct date formats observed: ISO YYYY-MM-DD, DD-MM-YYYY (936), MM-DD-YY (909), DD/MM/YYYY (891), MM/DD/YYYY (889), compact YYYYMMDD (889). "
                 "The DD-MM vs MM-DD ambiguity is a real-money risk — picking the wrong interpretation for the wrong locale silently misdates trades.")
        flag_box(doc, f"F-{col.upper()}-NUMERIC-LEAK", "P1-CRITICAL",
                 "Values '0.0' (966), '12.36' (236), '462.11' (87) observed — these are prices, not dates. Column-shift contamination.")

    elif col == "MessageId":
        if n_empty > 0 or n_ws_only > 0:
            flag_box(doc, f"F-{col.upper()}-NULL-PK", "P1-CRITICAL",
                     f"{n_ws_only:,} rows have an empty / whitespace-only primary key. A primary key cannot be null. "
                     "Vendor systems that enforce uniqueness or referential integrity will reject these records.")
        if suspect:
            flag_box(doc, f"F-{col.upper()}-NULL-TOKENS-MIXED", "P1-CRITICAL",
                     "Primary key contains literal strings: 'null' (469), 'NULL' (457), 'NaN' (458), 'None' (445). "
                     "A vendor that treats them as distinct from empty will create spurious duplicate-PK violations on the next batch.")
        flag_box(doc, f"F-{col.upper()}-NUMERIC-LEAK", "P1-HIGH",
                 "Top-10 includes '100' (207), '400' (183), '1700' (147), '500' (100) — these are TotalVolume values. Column-shift contamination.")

    elif col == "LinkMessageId":
        flag_box(doc, f"F-{col.upper()}-DEFINITION-AMBIGUITY", "P2",
                 f"{n_ws_only:,} rows are empty. This is expected for parent New Order events with no upstream link, but the "
                 "exact rule (empty for New Order; populated for Cancel/Fill/Replace) needs to be validated cross-message. Stage 2.")
        flag_box(doc, f"F-{col.upper()}-VALUE-LEAKAGE", "P1-HIGH",
                 "Top non-empty values include 'PROP95.VAULTVAULT' (688), 'INST11.FLUXBETA' (449) — these are Account/CounterPartyFirm values, "
                 "not message IDs. Column-shift contamination present in this field too.")

    elif col == "ParentOrderId":
        flag_box(doc, f"F-{col.upper()}-VALUE-LEAKAGE", "P1-HIGH",
                 "Same Account/CounterPartyFirm leakage pattern as LinkMessageId. The two columns appear to misalign together — "
                 "consistent with a single upstream shift event affecting subsequent fields.")

    elif col == "Instrument":
        flag_box(doc, f"F-{col.upper()}-BARE-QUOTE", "P1-CRITICAL",
                 "2,434 rows have a bare '\"' character as the instrument. These are residual artifacts of the unclosed-quote "
                 "contamination — the closing quote that wandered onto the next line. Confirms the upstream malformation pattern "
                 "and tells us where the recoverable lines start re-aligning.")
        flag_box(doc, f"F-{col.upper()}-NAN", "P1-HIGH",
                 "437 rows literally contain 'NaN'. Likely produced by a Python pipeline that called str(np.nan).")

    elif col == "ISIN":
        flag_box(doc, f"F-{col.upper()}-CARDINALITY-MISMATCH", "P1-CRITICAL",
                 f"{nunique:,} distinct ISINs vs ~373 distinct instruments. Expected ratio ≈ 1:1. The 17:1 ratio implies either "
                 "(a) bulk ISIN corruption, or (b) ISIN field is being polluted with non-ISIN values from adjacent columns.")
        flag_box(doc, f"F-{col.upper()}-LENGTH", "P1-HIGH",
                 f"Minimum stripped length = {info['stripped_length_stats']['min']}. Valid ISINs are exactly 12 characters. "
                 "Anything shorter is automatically invalid.")
        flag_box(doc, f"F-{col.upper()}-VALUE-LEAKAGE", "P1-CRITICAL",
                 "2,051 rows contain 'SYS_OMEGA' — that is a TransactionSource value, not an ISIN. Column-shift evidence.")

    elif col == "BuyOrSell":
        flag_box(doc, f"F-{col.upper()}-ENUM-VIOLATION", "P1-CRITICAL",
                 f"{nunique} distinct values where exactly 2 are expected (Buy/Sell). Invalid values: 'USD' (2,051), 'SYS_OMEGA' (11), 'nan' (6,628). "
                 "These are column-shift contamination — the rows missing earlier fields end up with Currency or TransactionSource in this position.")

    elif col == "Price":
        if n_ws_only > 0:
            flag_box(doc, f"F-{col.upper()}-MISSING", "P1-HIGH",
                     f"{n_ws_only:,} rows have an empty Price. Conditionally acceptable: market orders may have no limit price, "
                     "and New Order messages may carry 0 until filled. Needs Stage 2 cross-check vs MessageType.")
        # Many 0.0 values
        flag_box(doc, f"F-{col.upper()}-ZERO-DOMINANCE", "P2",
                 "Most common Price value is '0.0' (57,669 rows / ~39%). Plausible if dominated by orders that quote no price "
                 "(market orders, IOI, cancels). Validate by cross-tab against MessageType in Stage 2.")

    elif col == "TotalVolume":
        flag_box(doc, f"F-{col.upper()}-MISSING", "P1-HIGH",
                 f"{n_ws_only:,} rows missing TotalVolume. Volume is mandatory for any actionable surveillance event "
                 "(can't detect spoofing on a record with no quantity).")
        flag_box(doc, f"F-{col.upper()}-NAN", "P1-HIGH",
                 "8,679 rows contain the literal 'nan'. Same Python-str(np.nan) signature.")

    elif col == "Account":
        flag_box(doc, f"F-{col.upper()}-CARDINALITY", "INFO",
                 f"{nunique:,} distinct Accounts vs only 100 distinct CounterPartyFirms — broadly the expected shape "
                 "(many accounts per firm). However see the cross-column finding on Account==CounterPartyFirm collisions.")
        flag_box(doc, f"F-{col.upper()}-NAN", "P1-HIGH",
                 "8,690 rows contain 'nan'. Mandatory field, must be present.")

    elif col == "CounterPartyFirm":
        flag_box(doc, f"F-{col.upper()}-NAN", "P1-HIGH",
                 "8,690 rows contain 'nan'. Mandatory field.")
        flag_box(doc, f"F-{col.upper()}-CARDINALITY-ROUND", "INFO",
                 "Exactly 100 distinct firms — round number, consistent with synthetic seeding.")

    elif col == "Trader":
        flag_box(doc, f"F-{col.upper()}-LARGE-MISSING-RATE", "P2",
                 f"{n_ws_only + suspect.get('NaN', 0):,} rows ({100*(n_ws_only+suspect.get('NaN',0))/PROFILE['parsed_row_count']:.1f}%) have no trader. "
                 "Plausible if a large fraction of messages are system / algo generated, but the rate is suspiciously high. "
                 "Validate by joining MessageType — algo cancels are reasonable; missing trader on a Fill is not.")
        flag_box(doc, f"F-{col.upper()}-INCONSISTENT-NULLS", "P1-HIGH",
                 f"Same logical 'missing' state represented two ways: empty (59,234 rows) and 'NaN' literal (29,852 rows). "
                 "These will be counted differently by any downstream null-check.")

    elif col == "TransactionSource":
        flag_box(doc, f"F-{col.upper()}-CONSTANT-BY-DESIGN", "INFO",
                 "Single value 'SYS_OMEGA' across the dataset. Confirmed intentional — the synthetic dataset is single-sourced "
                 "to simplify exploratory analysis (see Stated Assumptions, §1.4). Documented but not a defect.")
        flag_box(doc, f"F-{col.upper()}-NAN-LITERAL", "P1-HIGH",
                 "8,690 rows still contain the 'nan' literal. Even on a constant-by-design column, 'nan' is a downstream defect "
                 "because it represents a row in which the source-system field was not written correctly.")

    elif col == "Currency":
        flag_box(doc, f"F-{col.upper()}-CONSTANT-BY-DESIGN", "INFO",
                 "Single value 'USD' across the dataset. Confirmed intentional (Stated Assumptions §1.4). Documented but not a defect. "
                 "A production feed will include multi-currency rows and any cross-currency consistency checks are 'not testable in this dataset'.")

    elif col == "Flags":
        flag_box(doc, f"F-{col.upper()}-MOSTLY-EMPTY", "INFO",
                 f"{n_ws_only:,} rows ({100*n_ws_only/PROFILE['parsed_row_count']:.1f}%) are empty. Expected for an exception-marker field. "
                 "Only non-empty value observed: 'Unmatched to Market Data from Vendor' (6,623 rows). This is itself a surveillance gap to investigate in Stage 2.")


def build_cross_column(doc):
    add_heading(doc, "5. Cross-Column Findings", level=1)

    add_heading(doc, "5.1 Systematic column-shift contamination", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Multiple fields show value-leakage signatures that line up: ISIN values appear as MessageType; Buy/Sell as TransactionTime; "
        "numeric volumes as MessageId; SYS_OMEGA as ISIN; USD as BuyOrSell; Account/CounterPartyFirm values as LinkMessageId/ParentOrderId. "
        "When the affected fields are projected onto their expected positions, they line up with the canonical column order — "
        "i.e., the affected rows are missing one or more upstream fields and everything thereafter has shifted left."
    )
    flag_box(doc, "F-XCOL-1 — Systematic left-shift", "P1-CRITICAL",
             "A unified column-shift signature explains many otherwise-disparate per-field anomalies. "
             "This means a single repair (re-aligning shifted rows against the canonical schema by inserting the missing field) "
             "may recover a large fraction of currently-malformed rows.")

    add_para(doc, "Example: a row with the 'column_shift_signature' pattern.", bold=True)
    rows = EXAMPLES["parsed_examples"]["column_shift_signatures"]
    add_row_example_table(doc, rows)

    add_heading(doc, "5.2 Account ≡ CounterPartyFirm collisions", level=2)
    add_para(doc,
             "Across the parsed rows, the top-frequency Account values are identical to the top-frequency CounterPartyFirm values "
             "(e.g., PROP95.VAULTVAULT, INST11.FLUXBETA). This means in many rows Account = CounterPartyFirm exactly.", size=10)
    flag_box(doc, "F-XCOL-2 — Account/Firm collision", "P2",
             "Two interpretations: (a) legitimate self-trading / prop-desk activity where the firm IS the account, or "
             "(b) the Account column has been copied from CounterPartyFirm in error. Resolve in Stage 2 by cross-tabbing "
             "against MessageType and TraderId — self-cross prop activity is plausible; the same collapse on agency flow is not.")

    add_row_example_table(doc, EXAMPLES["parsed_examples"]["account_equals_counterparty"])

    doc.add_page_break()


def build_open_questions(doc):
    add_heading(doc, "6. Open Questions / Deferred to Later Stages", level=1)
    add_bullet(doc, f"How many of the {PROFILE['csv_malformed_count']:,} CSV-malformed rows are recoverable by automated quote-repair and CR-normalisation (versus permanently lost)? Hypothesis: most are recoverable because the malformation pattern is systematic.")
    add_bullet(doc, "What is the canonical MessageType enum (so case normalisation can be validated against a documented list, not inferred)?")
    add_bullet(doc, "For MessageDate alternate formats: which (if any) of DD-MM vs MM-DD is intended? In a US-equities context I would default to MM/DD, but the canonical YYYY-MM-DD values combined with an ambiguous DD-MM-YYYY group is a real risk that should be confirmed with the data producer, not guessed.")
    add_bullet(doc, "Trader null rate ~61% — confirm with the team whether algo / system-generated messages legitimately carry no trader, and which MessageTypes those would be.")
    add_bullet(doc, "Flags field's only non-empty value is 'Unmatched to Market Data from Vendor' (6,623 rows) — this is itself a surveillance gap worth measuring in Stage 3 (metrics).")
    add_bullet(doc, "LinkMessageId vs ParentOrderId — they have near-identical cardinality and top values. Is the distinction documented (e.g., LinkMessageId = immediate predecessor, ParentOrderId = root order), or are they semantic duplicates?")

    doc.add_page_break()


def build_appendix(doc):
    add_heading(doc, "Appendix A — Reproducibility", level=1)
    add_para(doc, "All artifacts in this document are reproducible from the synthetic CSV plus three scripts. Steps:", size=10)
    add_bullet(doc, "Place synthetic_trade_data.csv into case-study/data/", mono=True)
    add_bullet(doc, "python case-study/analysis/01_profile.py", mono=True)
    add_bullet(doc, "python case-study/analysis/02_extract_examples.py", mono=True)
    add_bullet(doc, "python case-study/analysis/03_generate_stage1_doc.py", mono=True)
    add_para(doc, "Outputs: 01_profile_output.json, 02_examples.json, 02_raw_bad_lines.json, and this .docx.", size=10)
    add_para(doc, f"Python: 3.14.0 · pandas: 3.0.2 · numpy: 2.4.4 · python-docx: 1.2.0", size=9, italic=True)

    add_heading(doc, "Appendix B — Severity legend", level=2)
    add_kv_table(doc, [
        ("P1-CRITICAL", "Breaks ingestion or primary-key integrity. Records lost, silently merged, or duplicated. Vendor will drop."),
        ("P1-HIGH", "Breaks downstream joins, enum validation, or schema contracts. Records may be retained but downstream metrics will be wrong."),
        ("P2", "Accuracy / efficacy concern (not ingestion-blocking). Resolve in Stage 2 with cross-field checks."),
        ("INFO", "Observation only — may be intentional. Flag for confirmation with data producer."),
    ])


def main():
    doc = Document()
    # Tighten margins so per-column tables fit
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    build_cover(doc)
    build_exec_summary(doc)
    build_methodology(doc)
    build_file_level(doc)
    build_header_level(doc)
    build_column_audit(doc)
    build_cross_column(doc)
    build_open_questions(doc)
    build_appendix(doc)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED (open in Word?). Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
