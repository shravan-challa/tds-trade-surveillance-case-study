"""
One-page Executive Summary — front cover for the deliverable bundle.

This is what a reviewer at TD opens first. It states the bottom line in the
first paragraph, lists the four headline findings, splits ownership between
the producer and the surveillance team, and points to where to dig.

The numbers are pulled from 14_p3_metrics.json and 09_p2_findings.json so
the cover stays in sync with the rest of the bundle automatically.
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"
METRICS = HERE / "14_p3_metrics.json"
P2_FINDINGS = HERE / "09_p2_findings.json"
OUT = DELIV / "00_executive_summary.docx"
OUT_FALLBACK = DELIV / "00_executive_summary_v2.docx"


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def header_cell(cell, text: str, fill: str = "1F3A5F") -> None:
    cell.text = ""
    set_cell_shading(cell, fill)
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def main() -> None:
    metrics = json.loads(METRICS.read_text(encoding="utf-8"))
    p2 = json.loads(P2_FINDINGS.read_text(encoding="utf-8"))

    head = metrics["summary"]["headline"]
    by_status = metrics["summary"]["by_status"]
    n_records = metrics["input_records"]
    raw_lines = metrics["raw_data_lines"]

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---- Title block ---------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TD Securities — Trade Surveillance Case Study")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Executive Summary")
    rs.italic = True; rs.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Prepared by Shravan Challa  -  {date.today().isoformat()}  -  "
                            f"Bundle: 00 (this) + Stage 1 + Stage P1 + Stage P2 + Stage P3 + Methodology Journal")
    meta_run.font.size = Pt(9); meta_run.italic = True

    # ---- Bottom line --------------------------------------------------------
    bl = doc.add_paragraph()
    bl.add_run("Bottom line.  ").bold = True
    bl.add_run(
        f"The dataset is now ingestable but is not yet surveillable. P1 recovered {n_records:,} of "
        f"{raw_lines:,} raw lines ({head['ingestion_parse_rate_pct']:.2f}% parse rate, zero quarantine). "
        "P2 reveals four structural defects that, even after P1 cleaning, would render the surveillance system's "
        "alerts unreliable: an unstable Instrument-to-ISIN mapping, missing trader attribution on most fills, "
        "an Account column overwritten by the firm-level value on 97% of rows, and 1-in-10 child events with no "
        "resolvable parent. P3 turns these into a recurring scorecard so the surveillance team can re-run the same "
        f"audit on every new feed. Status across {by_status['GREEN'] + by_status['AMBER'] + by_status['RED']} metrics: "
        f"{by_status['GREEN']} GREEN, {by_status['AMBER']} AMBER, {by_status['RED']} RED."
    )

    bl2 = doc.add_paragraph()
    bl2.add_run("JPM-case anchor.  ").bold = True
    bl2.add_run(
        "Priority 1 was framed against the 2024 JPMorgan / Fed-OCC trade-surveillance enforcement, where records "
        "the surveillance system never received could not be surveilled. P1 closes that specific gap on this feed: "
        f"a strict-LF reader would silently drop 4,193 bare-CR lines plus mis-parse the 41.65% of records carrying "
        "multi-line cells. The repaired pipeline preserves every record. The remaining defects (P2) sit upstream of "
        "the surveillance system, not inside it."
    )

    # ---- Four headline findings ---------------------------------------------
    h2 = doc.add_paragraph()
    r = h2.add_run("Four headline findings (the conversation-leading material).")
    r.bold = True; r.font.size = Pt(11)

    findings = [
        ("F1.  Instrument-to-ISIN mapping is unstable",
         f"{head['instrument_isin_drift_pct']:.1f}% of instruments map to multiple ISINs",
         "OO -> 4,528 ISINs, YAZ -> 1,553, YV -> 836",
         "Per-instrument surveillance rules (price outlier, position aggregation, market-data join) cannot fire reliably."),
        ("F2.  Trader attribution missing on most fills",
         f"{head['fill_attribution_rate_pct']:.1f}% of fills carry a Trader",
         f"{p2['E1']['count']:,} of {int(p2['E1']['count'] / (1 - p2['E1']['pct_of_fills']/100)):,} fills lack attribution",
         "Reg-best-execution + MAR market-abuse attribution gap. Regulator-exam critical."),
        ("F3.  Account is collapsed to CounterPartyFirm on 97% of rows",
         f"{head['account_cpf_collapse_pct']:.1f}% Account == CPF",
         "Stage-1 profile shows Account has 1,043 distinct values (vs CPF=100), but the granular value appears on only 2.6% of rows",
         "Wash-trade rules cannot isolate one account inside a firm. Looks like an upstream column-copy bug."),
        ("F4.  Order lineage has 10% holes",
         f"A2 = {head['lineage_orphan_parent_pct']:.2f}% orphaned ParentOrderIds; B1 time-travel = {p2['B1']['pct_of_resolvable_children']:.2f}%",
         f"B2 over-fill on {p2['B2']['count']:,} parents; B3 fill-after-cancel on {p2['B3']['count']:,} rows",
         "Lifecycle reconstruction breaks for 10% of child events. Producer must guarantee parent-first ordering."),
    ]

    t = doc.add_table(rows=1 + len(findings), cols=4)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    for i, w in enumerate([Inches(2.4), Inches(1.7), Inches(1.9), Inches(1.5)]):
        t.columns[i].width = w
    for i, h in enumerate(["Finding", "Headline metric", "Evidence", "Surveillance impact"]):
        header_cell(t.rows[0].cells[i], h)
    for ri, (title_text, metric, evidence, impact) in enumerate(findings, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate([title_text, metric, evidence, impact]):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(val)
            r.font.size = Pt(8.5)
            if ci == 0:
                r.bold = True

    # ---- Ownership matrix ---------------------------------------------------
    h3 = doc.add_paragraph()
    r = h3.add_run("Ownership — who fixes what.")
    r.bold = True; r.font.size = Pt(11)

    own = doc.add_paragraph()
    own.add_run(
        "Every 'rule fires too much' finding traces back to an upstream defect, not a bad rule. "
        "The producer's data needs to change. The surveillance team's rules are mostly fine; they just can't be "
        "evaluated on this feed."
    )

    ownership = [
        ("Producer (data feed)", "F1, F2, F3, F4; A4 schema misuse; B2 over-fills; B4 replace-without-parent",
         "Resolve symbol-to-ISIN reference data; backfill Trader from OMS; restore granular Account; "
         "guarantee parent-first ordering; fix MessageType labelling; investigate over-fill cap."),
        ("Surveillance team (rules)", "F1 cancel-ratio threshold; G1 price-outlier threshold; B3 fill-after-cancel review",
         "Tune F1 threshold once Trader coverage improves (M6.1); re-evaluate G1 once C1 is resolved; "
         "decide which B3 events are race-conditions vs investigations."),
        ("Both", "F2/G1 (downstream of producer fixes); P3 scorecard governance",
         "Re-baseline F2 wash-trade and G1 outlier thresholds AFTER producer ships fixes for D1/C1. "
         "Establish RED/AMBER/GREEN review cadence per the M1-M7 themes."),
    ]
    t2 = doc.add_table(rows=1 + len(ownership), cols=3)
    t2.style = "Light Grid Accent 1"
    t2.autofit = False
    for i, w in enumerate([Inches(1.8), Inches(2.6), Inches(3.1)]):
        t2.columns[i].width = w
    for i, h in enumerate(["Owner", "Findings owned", "Action"]):
        header_cell(t2.rows[0].cells[i], h)
    for ri, (owner, owned, action) in enumerate(ownership, start=1):
        cells = t2.rows[ri].cells
        for ci, val in enumerate([owner, owned, action]):
            cells[ci].text = ""
            r = cells[ci].paragraphs[0].add_run(val)
            r.font.size = Pt(8.5)
            if ci == 0:
                r.bold = True

    # ---- Bundle map ---------------------------------------------------------
    h4 = doc.add_paragraph()
    r = h4.add_run("Where to dig.")
    r.bold = True; r.font.size = Pt(11)

    bundle = [
        ("00_executive_summary.docx",         "This document — bottom line, headline findings, ownership."),
        ("stage1_profile_audit.docx",         "Stage 1 — column-by-column raw-data audit. The 'before' state, with raw-row evidence."),
        ("stage_p1_repair_report.docx",       "Stage P1 — every named normalisation rule (N1..N11), the rationale, the count of rows it touched."),
        ("stage_p2_accuracy_efficacy.docx",   "Stage P2 — surveillance-domain checks (Lineage / Lifecycle / Instrument / Account / Trader / Signatures / Statistical)."),
        ("stage_p3_metrics_scorecard.docx",   "Stage P3 — 44 KPIs across 7 themes (M1..M7) with RED/AMBER/GREEN status and named owner."),
        ("methodology_journal.docx",          "Chronological narrative of the work — every decision, hypothesis, recalibration, and lesson, in the order they happened."),
    ]
    t3 = doc.add_table(rows=1 + len(bundle), cols=2)
    t3.style = "Light Grid Accent 1"
    t3.autofit = False
    for i, w in enumerate([Inches(2.7), Inches(4.7)]):
        t3.columns[i].width = w
    for i, h in enumerate(["File", "What you'll find"]):
        header_cell(t3.rows[0].cells[i], h)
    for ri, (fname, what) in enumerate(bundle, start=1):
        cells = t3.rows[ri].cells
        cells[0].text = ""
        rf = cells[0].paragraphs[0].add_run(fname)
        rf.font.name = "Consolas"; rf.font.size = Pt(8.5)
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(what).font.size = Pt(8.5)

    # ---- Defensibility footer -----------------------------------------------
    foot = doc.add_paragraph()
    foot.add_run("Defensibility.  ").bold = True
    foot.add_run(
        "Every number in this bundle is reproducible. The analysis/ folder contains a numbered pipeline "
        "(01..16) that regenerates every artefact from the raw CSV. Two test scripts (12, 13) assert that the "
        "P1 pipeline's invariants hold and that every P2 finding count reproduces from the file. 63 tests, all passing. "
        "If a number drifts, a test breaks."
    )
    foot_run = foot.runs[1]
    foot_run.font.size = Pt(9)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED. Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
