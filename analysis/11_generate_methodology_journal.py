"""
Methodology Journal — chronological narrative of how the work proceeded,
the decisions made at each step, the hypotheses chased, the moments of doubt,
and the recalibrations. Companion to the three formal Stage docs.

The audience for this document is the interviewer asking 'walk me through your
thinking'. It explicitly captures the reasoning behind each decision so a reader
can interrogate any choice without needing to re-derive it from the code.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"
OUT = DELIV / "methodology_journal.docx"
OUT_FALLBACK = DELIV / "methodology_journal_v2.docx"


# -------- docx primitives -----------------------------------------------------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def p(doc, text, *, bold=False, italic=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return para


def bullet(doc, text, *, mono=False):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    if mono:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    else:
        run.font.size = Pt(10)
    return para


def quote_block(doc, text):
    """Indented italicised block for direct quotes from the brief / from data / from raw output."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.right_indent = Inches(0.5)
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    return para


def code_block(doc, text):
    """Monospace code-like block for command/raw-output snippets."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    return para


def kv_table(doc, pairs, c0=2.0, c1=4.5):
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    t.columns[0].width = Inches(c0)
    t.columns[1].width = Inches(c1)
    for i, (k, v) in enumerate(pairs):
        a, b = t.rows[i].cells[0], t.rows[i].cells[1]
        a.text = ""; b.text = ""
        ra = a.paragraphs[0].add_run(str(k)); ra.bold = True; ra.font.size = Pt(9)
        rb = b.paragraphs[0].add_run(str(v)); rb.font.size = Pt(9)


def callout(doc, label, kind, text):
    """Coloured callout. kind: DECISION | HYPOTHESIS | INSIGHT | RECALIBRATION | LESSON."""
    color = {
        "DECISION": "2980B9",
        "HYPOTHESIS": "8E44AD",
        "INSIGHT": "27AE60",
        "RECALIBRATION": "E67E22",
        "LESSON": "16A085",
    }.get(kind, "7F8C8D")
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.2)
    t.columns[1].width = Inches(5.3)
    a, b = t.rows[0].cells[0], t.rows[0].cells[1]
    a.text = ""; b.text = ""
    set_cell_shading(a, color)
    set_cell_shading(b, "F4F6F8")
    ra = a.paragraphs[0].add_run(kind); ra.bold = True; ra.font.size = Pt(8)
    ra.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp = b.paragraphs[0]
    r1 = pp.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(9)
    r2 = pp.add_run(text); r2.font.size = Pt(9)


# -------- sections ------------------------------------------------------------
def cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TD Securities — Trade Surveillance Case Study")
    r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Methodology Journal")
    rs.italic = True; rs.font.size = Pt(14)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Chronological narrative of the work, the decisions, and the train of thought")
    rs2.italic = True; rs2.font.size = Pt(11)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: Shravan Challa\n").font.size = Pt(10)
    meta.add_run(f"Date: {date.today().isoformat()}\n").font.size = Pt(10)
    meta.add_run("Companion to: Stage 1 Audit Trail, Stage P1 Repair Report, Stage P2 Accuracy & Efficacy Report\n").font.size = Pt(10)

    doc.add_paragraph()
    pp = doc.add_paragraph()
    pp.add_run("Purpose. ").bold = True
    pp.add_run(
        "The three Stage reports describe WHAT was found and WHAT was done. This document describes WHY each step "
        "happened in the order it did, what hypothesis I was chasing, and where I had to recalibrate. The intended "
        "audience is the interviewer asking 'walk me through your thinking' — every decision below is named and "
        "rationalised so it can be defended in conversation."
    )
    doc.add_page_break()


def opening_frame(doc):
    add_heading(doc, "1. Opening Frame — What the Brief Told Me, and What I Decided It Meant", level=1)
    p(doc, "The brief gave three ranked priorities and one explicit anchor:")
    quote_block(doc,
                'Priority 1 is ensuring the dataset is free of errors and has the correct format for ingestion to ensure '
                'the vendor does not drop the records which would impact completeness (JPM Case as you mentioned). '
                'Priority 2 is accuracy and efficacy of data. Priority 3 is generation of metrics.')
    p(doc, "Three things from the brief shaped how I worked:")
    bullet(doc, "JPM Case is named explicitly. That's the user's anchor — a 2024 surveillance-completeness enforcement where records the surveillance system never received could not be surveilled. Priority 1 is therefore not 'tidy CSV', it is 'don't let the vendor silently drop records'. I made completeness the headline metric.")
    bullet(doc, "Priority 2 is open-ended ('based on your imagination'). That's the differentiator section — it's where domain framing matters more than tooling. I decided the lens would be surveillance-specific, not generic DQ (more on this in §10).")
    bullet(doc, "'Defensible' is the keyword. Every choice must have a stated rationale and be reproducible from the code, not from memory. This shaped the artifact structure (every script has a stable ID; every rule has a stable name; every count comes from JSON I can re-derive).")

    callout(doc, "Opening decision — measure completeness, not tidiness", "DECISION",
            "The primary metric is 'records that survive ingestion as a percentage of records in the file'. Whitespace, "
            "null tokens, case variation are real defects but secondary; record loss is JPM-class and primary.")
    doc.add_page_break()


def first_contact(doc):
    add_heading(doc, "2. First Contact With the Data (no parsing yet)", level=1)
    p(doc, "Before writing any analysis I ran two near-zero-cost commands. They both produced a finding.")

    p(doc, "Command 1 — peek at structure:", bold=True)
    code_block(doc, "head -3 synthetic_trade_data.csv")
    p(doc, "Three observations from the output alone:", italic=True, size=9)
    bullet(doc, "Every field shows ' , ' (space-comma-space) separators. The whitespace is on every column AND every value, including the header row.")
    bullet(doc, 'Row 2 already contains the literal string "NaN" in the Trader column — and an empty " " value in another field. Two different null representations in the first two rows of data.')
    bullet(doc, "MessageId is composite: YYYY-MM-DD.RTnnn.X.nnn. TransactionTime is ISO 8601 with millisecond precision. These are the strongest type hints in the file.")

    p(doc, "Command 2 — count rows:", bold=True)
    code_block(doc, "wc -l synthetic_trade_data.csv\n246522 …")
    p(doc, "Filed away as the raw baseline. (Spoiler — this number was wrong, and the way I caught the error became one of the better findings in the report. See §5.)", italic=True, size=9)

    callout(doc, "First-contact hypothesis set", "HYPOTHESIS",
            "Before writing code I had three live hypotheses: (1) pervasive whitespace will defeat exact-match joins and "
            "schema validation; (2) mixed null tokens will confuse downstream nullability checks; (3) the dataset spans "
            "multiple physical conventions that need to be normalised before any semantic check is meaningful.")
    doc.add_page_break()


def first_profile(doc):
    add_heading(doc, "3. The First Profile — and Why It Failed", level=1)
    p(doc, "Wrote 01_profile.py. The configuration choices were deliberate and the reasoning matters:")

    kv_table(doc, [
        ("dtype=str",
         "Read everything as string. If I let pandas infer types, '0.0' silently becomes 0.0 (float) and the dataset's "
         "type-mismatch defects vanish from the profile. To detect defects I need to see the bytes."),
        ("keep_default_na=False, na_filter=False",
         "Disable pandas' built-in null inference. Otherwise 'NaN', 'NULL', 'None' all collapse to a single np.nan and "
         "I can no longer count how many of each token appear (or that they're inconsistent across the same column)."),
        ("Universal-newline file read",
         "Match what Power Query, Excel, and modern parsers would do. This bit me later — see §5."),
    ])

    p(doc, "First run terminated with a parser error:")
    code_block(doc,
               "pandas.errors.ParserError: Error tokenizing data. C error: Expected 18 fields in line 40, saw 19")

    callout(doc, "The parser error is itself the first P1 finding", "INSIGHT",
            "The fact that the C parser exploded on line 40 is more useful than the count it would have produced. "
            "It tells me there are records the vendor system cannot tokenize at all — the JPM-class scenario.")

    p(doc, "Looked at lines 38–42:")
    code_block(doc,
               'Line 42: XFKA , New Order , 2026-02-05T08:00:14.570Z , … , "KRJS\n'
               'Line 43:  , USKGV1OINAL2 , Buy , 0.0 , 100 , …')

    callout(doc, "Unclosed-quote hypothesis", "HYPOTHESIS",
            'Some Instrument fields are quoted multi-line cells, but the producer split the value with a line break — '
            'so "KRJS opens on line 42 and the closing " is on a later line. The CSV parser, seeing a quoted field that '
            'never closes, absorbs subsequent lines into one record and produces a wrong field-count error.')

    p(doc, "Switched the parser to tolerant mode (engine='python', on_bad_lines='skip') so I could measure the malformation rate rather than abort:")
    code_block(doc, "Parsed: 146,304 rows x 18 cols\nCSV-malformed (skipped or merged): 100,217 rows (40.65%)")
    callout(doc, "Stage 1 headline number", "INSIGHT",
            "40.65% of records would be silently dropped by a strict ingestion pipeline. This is the number the entire "
            "deliverable is anchored on. Everything else is sub-headline.")
    doc.add_page_break()


def evidence_extraction(doc):
    add_heading(doc, "4. Building the Evidence Pack (Stage 1 Audit Trail)", level=1)
    p(doc,
      "At this point I had aggregate counts but no row-level evidence. Aggregate counts are easy to dismiss "
      "(\"how do you know it's really unclosed quotes and not something else?\"). Wrote 02_extract_examples.py to pull "
      "concrete row examples per defect class. Each defect class got 3–5 row examples; each ambiguous case got a "
      "raw-file line number so a reader can open the CSV and verify themselves.")

    p(doc, "Then wrote 03_generate_stage1_doc.py to assemble the audit document. Some choices in that doc deserve to be named:")
    bullet(doc, "Each column got its own section: Role (what the field is for) → Expected (what valid values look like) → Observed (what is in the file) → Flags (what's wrong) → Row-level examples → Reasoning.")
    bullet(doc, "Severity tags (P1-CRITICAL, P1-HIGH, P2, INFO) make the document scannable by a non-technical reader.")
    bullet(doc, "Every flag has a stable ID (e.g., F-FILE-1, F-MESSAGEID-NULL-PK) so the doc and the code can reference the same finding.")

    callout(doc, "Why this format instead of a notebook", "DECISION",
            "An interviewer who is not at a terminal needs to be able to read the work end-to-end. A docx with column-level "
            "tables and severity-coded flag boxes carries the work as well as a Jupyter notebook would — and the docx is "
            "portable, sharable, and easier to discuss line-by-line in conversation.")
    doc.add_page_break()


def power_query_moment(doc):
    add_heading(doc, "5. The Power Query Reconciliation Moment", level=1)
    p(doc, "You flagged a discrepancy I would not have caught on my own:")
    quote_block(doc, "Loading the dataset using power query shows 250715 rows (including headers). Am i missing something?")

    callout(doc, "Initial reaction", "RECALIBRATION",
            "Two numbers, one truth. Don't pick one tool and dismiss the other — find out which is right and why. "
            "Either Power Query is inflating the count, OR my wc -l undercounted, OR there is a genuine ambiguity in "
            "what counts as a 'row' in this file.")

    p(doc, "Wrote 04_line_count_audit.py to count line-terminator characters at the byte level — without invoking any CSV parser:")
    code_block(doc,
               "Byte-level newline counts:\n"
               "  \\n total:           246,522\n"
               "  \\r total:           250,715\n"
               "  \\r\\n pairs:         246,522\n"
               "  bare \\r:           4,193\n"
               "  bare \\n:           0\n"
               "\n"
               "Universal-newline line count (Python text mode): 250,715")

    callout(doc, "Diagnosis", "INSIGHT",
            "The file has 246,522 CRLF terminators AND 4,193 bare-CR terminators. Power Query and Python text mode use "
            "universal-newline reading and see all 250,715 logical lines. Git Bash wc -l counts only \\n characters and "
            "misses the 4,193 bare-CR-terminated lines. Your number was right.")

    p(doc, "Implications:", bold=True)
    bullet(doc, "Updated the baseline to 250,714 data lines. The malformation rate became 41.65% (slightly worse than my initial 40.65%).")
    bullet(doc, "Mixed line endings is itself a P1-CRITICAL finding (a strict-LF reader would silently lose 4,193 records — the invisible kind of data loss).")
    bullet(doc, "Documented the reconciliation across tools in the Stage 1 audit so a future reader running Power Query, Excel, or Git Bash can verify the numbers themselves rather than trusting a single source.")

    callout(doc, "Lesson", "LESSON",
            "Pay attention when a user flags a mismatch — it is almost always a finding, even if it initially looks like noise. "
            "The cost of investigating an off-by-4193 row count was 30 minutes; the value was a P1-CRITICAL flag that strengthens the deliverable.")
    doc.add_page_break()


def stage_p1_thinking(doc):
    add_heading(doc, "6. Designing the P1 Pipeline — the Two Insights That Changed the Plan", level=1)

    add_heading(doc, "6.1 Reading the malformation pattern more carefully", level=2)
    p(doc,
      "Stage 1 surfaced 104,410 malformed rows (after the line-count correction). The naive plan was: classify them by failure mode, "
      "fix what is mechanical, quarantine what is not. I almost wrote that pipeline. Then I looked again at the bad-line examples:")
    code_block(doc,
               'Line 15: XFKB , New Order , … , "          (closing " not on this line)\n'
               'Line 16: LJ" , USUJGWJM9EI1 , Buy , …      (closing " here, then a full record)')

    callout(doc, "Key observation", "INSIGHT",
            "Every bad line followed the SAME structure: a quoted field that should contain a multi-line value, "
            "but the quote handling is defeated by something. If the cause is structural and uniform, the fix is "
            "structural and uniform too — not record-by-record.")

    add_heading(doc, "6.2 Why the quoting was failing — the RFC 4180 nuance", level=2)
    p(doc,
      "RFC 4180 (and Python's csv module by default) only recognises a field as 'quoted' if the opening quote is the FIRST "
      "character of the field. The synthetic data uses ' , ' (space-comma-space) as the separator. So a field that should look "
      'like ` "KRJS"` is actually ` "KRJS"` — and that leading space turns the opening quote into a literal character. The '
      "csv module then treats embedded line breaks inside the quoted value as record separators, not as part of the field.")

    callout(doc, "Decision — fix the cause, not the symptom", "DECISION",
            "Don't try to repair 104,410 individual broken rows. Tell the CSV parser to skip leading whitespace before "
            "the quote-detection state machine starts (csv.reader supports skipinitialspace=True). Normalise the bare-CR "
            "line endings in the same pass. Test whether the structural fix recovers the rows on its own.")

    add_heading(doc, "6.3 The result", level=2)
    p(doc, "Ran 05_repair_structural.py. The numbers were better than I expected:")
    kv_table(doc, [
        ("Stage 1 strict-parsed", "146,304 rows"),
        ("After structural repair", "242,429 rows"),
        ("Recovered", "+96,125 (92.06% of the 104,410 malformed)"),
        ("Quarantine after repair", "0 rows"),
        ("Final parse-rate vs raw lines", "96.70%"),
    ])

    callout(doc, "What I expected vs what happened", "INSIGHT",
            "I expected maybe 50% recovery and a quarantine bucket to investigate. The actual 92% recovery — with zero "
            "quarantine — meant the malformation was even more uniform than I thought. The remaining ~8,285 'missing' "
            "lines weren't lost; they were correctly absorbed as embedded content within parent records' multi-line cells.")
    doc.add_page_break()


def stage_p1_value_normalisation(doc):
    add_heading(doc, "7. P1 Value Normalisation — Two Recalibrations", level=1)

    p(doc, "Wrote 06_normalize_values.py. The structure I chose was deliberate:")
    bullet(doc, "Each transformation is a named rule (N1.ExchangeId, N2, N3a, etc.) with rule_id, column, rationale (the WHY), action (the WHAT), and rows-affected count.")
    bullet(doc, "The pipeline emits a per-row flag bitset (06_row_flags.tsv) so downstream stages know which rows were touched by which rule. This is the audit trail.")
    bullet(doc, "Conservative defaults: when in doubt, null + flag rather than guess. Reason: P1 must be defensible without producer review; aggressive recovery is P2.")

    add_heading(doc, "7.1 Recalibration #1 — the PK regex bug", level=2)
    p(doc, "First run reported 139,812 MessageIds (57.7%) as malformed. Initial reaction: 'really?' That number is too high to be plausible for a primary-key violation — if 57% of rows had a bad PK, the dataset would be unusable for any purpose.")

    callout(doc, "Don't trust large counts before sanity-checking them", "RECALIBRATION",
            "Wrote a small inline diagnostic to dump the actual MessageId 'shapes' (digits→#, letters→X). The top shapes were "
            "####-##-##.XX###.X.######, ####-##-##.XX###.X.#### , ####-##-##.XX###.X.#####. My regex hardcoded the middle "
            "segment as '.P.' but the data has '.[A-Z].' — at least 'P' and 'C' both occur.")

    p(doc, "Fixed the regex to [A-Z]; reran; the genuine-malformed count dropped to 0. Lesson:")
    callout(doc, "Lesson", "LESSON",
            "A surprisingly large count is more often a regex bug than a real defect. Always validate a regex against "
            "actual examples before reporting the count as a finding.")

    add_heading(doc, "7.2 Recalibration #2 — the 'column-shift was a parser artifact' insight", level=2)
    p(doc, "After the regex fix, several of my defect counts returned zero:")
    bullet(doc, "N3b (MessageDate unparseable): 0 rows")
    bullet(doc, "N4 (TransactionTime not ISO 8601): 0 rows")
    bullet(doc, "N6 (ISIN format invalid): 0 rows")
    bullet(doc, "N7 (BuyOrSell ∉ {Buy, Sell}): 0 rows")

    p(doc, "Initial reaction: 'are my checks broken?' The Stage 1 audit had documented thousands of these — ISINs appearing as MessageType, 'Buy'/'Sell' appearing as TransactionTime, 'SYS_OMEGA' appearing as ISIN. None of those signals should disappear by accident.")

    callout(doc, "The real explanation", "INSIGHT",
            "The 'column-shift contamination' diagnosed at Stage 1 was a downstream artifact of the unclosed-quote misparse — "
            "when csv.reader incorrectly split a multi-line quoted cell, the resulting records had fields in the wrong positions. "
            "Once skipinitialspace=True let the quote handling work, every value landed in its intended column. "
            "The structural repair (06) did far more semantic work than I realised; many 'data corruption' defects were never "
            "in the data — they were in the parser's output.")

    callout(doc, "Why this is worth highlighting", "DECISION",
            "Naming this finding explicitly in the P1 doc strengthens the deliverable: it shows the candidate can distinguish "
            "between data-corruption defects and parser-induced artifacts. A naïve report would have spent pages discussing "
            "'column shift' as a real defect class; the correct report retires it as a parser symptom.")

    add_heading(doc, "7.3 Recalibration #3 — the misleading 99% 'any flag' number", level=2)
    p(doc, "Post-P1 profile showed 99.12% of rows had at least one P1 flag fired. Initial reaction: 'did the cleaning fail?'")
    callout(doc, "Diagnosis", "RECALIBRATION",
            "The 99% number was including 'NULL_<column>' flags for columns where the empty value is the EXPECTED state — "
            "LinkMessageId / ParentOrderId are empty for parent New Orders by design; Trader is empty for algo flow; the "
            "Flags column is empty for non-exception messages. I was counting 'null token was unified' (mechanical fix) the "
            "same way I was counting 'price was negative' (real defect). Split the flag set into 'defect-indicative' and "
            "'legitimate-empty' in the P1 doc.")

    callout(doc, "Lesson", "LESSON",
            "Flags need policy. 'A flag fired on this row' is meaningless without knowing which flags reflect defects vs. "
            "which reflect optional fields being legitimately empty. Always classify flags before aggregating to a row-health score.")
    doc.add_page_break()


def stage_p2_strategy(doc):
    add_heading(doc, "8. P2 Strategy — Choosing the Lens", level=1)

    p(doc, "The brief was emphatic that P2 is open-ended:")
    quote_block(doc, "this is an open-ended question for priority 2, … based on your imagination.")

    p(doc, "Two ways to interpret this:")
    bullet(doc, "Generic DQ lens — completeness/accuracy/timeliness/consistency, applied to every column. Safe and shallow.")
    bullet(doc, "Surveillance lens — every check is framed as 'what surveillance rule does this support, and would a regulator be satisfied with the data?'. Higher bar; needs domain knowledge.")

    callout(doc, "Choosing the surveillance lens", "DECISION",
            "I chose the surveillance lens. The role is 'Trade Surveillance Data', not 'Data Quality Engineer'. The interviewer "
            "named the JPM case as the framing in Priority 1. Continuing that frame into Priority 2 is consistent and "
            "demonstrates I understand what the data is FOR, not just what the data IS.")

    p(doc, "Organised checks into seven themes. The themes are chosen to map onto the surveillance rules a regulator expects:", bold=True)
    kv_table(doc, [
        ("A. Referential integrity (lineage)", "Can surveillance reconstruct the order lifecycle?"),
        ("B. Order lifecycle coherence", "Are the events in a physically-possible sequence?"),
        ("C. Instrument / ISIN coherence", "Is the symbol-to-ISIN mapping stable enough for per-instrument rules?"),
        ("D. Account / Firm relationships", "Prop self-cross vs agency wash-trade — can the rules distinguish?"),
        ("E. Trader attribution", "Can every execution be tied to a human (or attested algo)?"),
        ("F. Surveillance signatures", "Cancel ratios, wash-trade timing — direct detection."),
        ("G. Statistical anomalies", "Per-instrument outliers — low-cost first-pass triage."),
    ])
    doc.add_page_break()


def stage_p2_execution(doc):
    add_heading(doc, "9. P2 Execution — One Bug, One Big Finding", level=1)

    add_heading(doc, "9.1 Bug — parent-lookup index had duplicates", level=2)
    p(doc, "Wrote 09_p2_accuracy_efficacy.py. First run failed:")
    code_block(doc, "pandas.errors.InvalidIndexError: Reindexing only valid with uniquely valued Index objects")
    p(doc, "Cause: the parent-timestamp lookup was indexed by MessageId, but we already knew from P1 (rule N5c) that 17,671 rows participate in PK duplicates. Series.map cannot reindex against a non-unique index.")
    callout(doc, "Fix and lesson", "LESSON",
            "drop_duplicates(subset='MessageId', keep='first') on the parent lookup. But also recorded the duplicate count "
            "(1,722 duplicate New-Order PKs collapsed) as an explicit input to the P2 results — the duplicate-PK problem "
            "from P1 is not orthogonal to P2; it shapes which checks are reliable.")

    add_heading(doc, "9.2 The findings that matter", level=2)
    p(doc, "Out of 14 checks, four findings stood out as 'lead-the-conversation' material:")
    kv_table(doc, [
        ("C1 — Instrument↔ISIN mapping unstable",
         "222 instruments map to multiple ISINs. ~60% of the instrument universe. Every per-instrument surveillance rule "
         "is unreliable until this is resolved. This is the headline DQ finding because it's specific, quantified, and "
         "domain-specific."),
        ("E1 — 65% of Fills have no Trader",
         "Direct Reg-best-execution / MAR-attribution gap. Easy to explain to a non-technical reader; impossible to "
         "ignore on a regulator exam."),
        ("B2 — Fill total > parent order volume on 1,748 parents",
         "Textbook execution-integrity check. Every surveillance team has a rule for this; every regulator looks for it. "
         "The signal is pure: no domain ambiguity, no threshold-tuning required."),
        ("D1b — Non-PROP self-cross on 62% of rows",
         "Either seeded baseline OR the Account column is a copy of CounterPartyFirm. The Stage-1 profile shows Account "
         "has 1,043 distinct values vs CPF's 100, so the granular Account values DO exist — they're just only on 3% of rows. "
         "The pattern smells like an upstream column-copy bug. Either way, surveillance can't isolate accounts inside a firm."),
    ])

    callout(doc, "What I deliberately left as 'needs tuning'", "DECISION",
            "F1 (cancel-to-new-order ratio per Trader) and G1 (price outliers via 3·IQR) used placeholder thresholds. I "
            "did NOT calibrate them to look impressive. Both rules are surveillance-real, but the thresholds belong to the "
            "surveillance team — not to me at the case-study stage. Saying 'placeholder, please tune' is more defensible "
            "than pretending I picked the right number.")
    doc.add_page_break()


def decision_register(doc):
    add_heading(doc, "10. Decision Register", level=1)
    p(doc, "Every non-obvious choice, named and rationalised for the interview conversation:", italic=True, size=9)

    decisions = [
        ("Read all columns as string at Stage 1",
         "Lets me see 'NaN', 'NULL', '0.0', '12.36' as the literal bytes they are, without pandas hiding them through type "
         "coercion. The defects ARE the byte representation; if I coerce types I lose the signal."),
        ("Disable pandas null inference at Stage 1",
         "Same reason — I want to count how many of each null-token variant exist, not see them all collapse to one np.nan."),
        ("Use Python's text-mode universal-newline reading for line counting",
         "Matches Power Query, Excel, modern parsers. Strict-LF (Unix wc -l) silently misses bare-CR lines. The Power Query "
         "reconciliation in §5 confirmed this was the right call."),
        ("engine='python' + on_bad_lines='skip' for the tolerant first parse",
         "I want to MEASURE the malformation rate, not abort on the first bad line. Cost: slower parser. Benefit: real number."),
        ("skipinitialspace=True for the P1 structural repair",
         "RFC 4180 only recognises quoted fields when the opening quote is the first character. The synthetic data has ' , ' "
         "separators that defeat that recognition. skipinitialspace=True consumes the leading space before quote detection. "
         "Recovered 92% of malformed rows with this one flag."),
        ("Named rules in the value-normalisation script (N1.ExchangeId, N2, N3a, …)",
         "Every transformation gets a stable ID. Every doc cross-reference cites the ID. If someone asks 'why did 4,105 ExchangeId "
         "values become null?' I can point to rule N1.ExchangeId and walk through the rationale, the token list, and the count."),
        ("Flag-don't-fix for identifier format checks (Account, Trader, ExchangeId)",
         "A pattern mismatch may indicate corruption, but it may also indicate legitimate format variation across desks. "
         "Conservative P1 surfaces the flag for review; aggressive P1 would null out values that are actually correct. "
         "When in doubt at the case-study stage, surface > delete."),
        ("Surveillance lens for P2 instead of generic DQ",
         "The role is a surveillance role. The brief named the JPM case. The lens has to be 'what surveillance rule does "
         "this data support?' rather than 'is this data tidy?'. Differentiates the response from a generic DQ analysis."),
        ("Drop duplicates on parent lookup before mapping",
         "Practical fix for the InvalidIndexError, but also recorded explicitly because the duplicate-PK problem from P1 "
         "constrains what P2 can confidently measure. Acknowledging the constraint is better than silently working around it."),
        ("Reproducible scripts over a Jupyter notebook",
         "Three reasons: (1) artifacts can be regenerated bit-for-bit without manual cell-running; (2) each script has a "
         "single responsibility and a single output, so the dependency graph is explicit; (3) the work can be code-reviewed "
         "the same way production code is."),
        ("Word .docx as the deliverable format",
         "Portable, sharable, scannable by a non-technical reader. Each report follows the same column-level / per-rule "
         "structure so a reader who has seen one can navigate the others without orientation cost."),
        ("Document open questions explicitly rather than guessing",
         "Multiple findings (the Account-CPF collapse, the 5,142 New Orders with parents, the DD-MM vs MM-DD date "
         "ambiguity) require producer confirmation. Listing them as open questions is more defensible than picking an "
         "interpretation and hoping it's right."),
    ]
    for label, why in decisions:
        callout(doc, label, "DECISION", why)
    doc.add_page_break()


def reproducibility_map(doc):
    add_heading(doc, "11. Reproducibility Map", level=1)
    p(doc, "Running the scripts in this order from case-study/analysis/ regenerates every artifact (TSV, JSON, docx) without manual intervention:", size=10)
    kv_table(doc, [
        ("01_profile.py",              "Stage 1 raw profile → 01_profile_output.json"),
        ("02_extract_examples.py",     "Stage 1 row-level evidence → 02_examples.json, 02_raw_bad_lines.json"),
        ("03_generate_stage1_doc.py",  "Stage 1 audit document → deliverable/stage1_profile_audit.docx"),
        ("04_line_count_audit.py",     "Line-ending diagnostic for the Power Query reconciliation"),
        ("05_repair_structural.py",    "P1 structural repair → cleaned_stage_p1.tsv + 05_repair_stats.json"),
        ("06_normalize_values.py",     "P1 value-level normalisation → cleaned_p1_final.tsv + 06_row_flags.tsv + 06_normalize_stats.json"),
        ("07_post_p1_profile.py",      "Post-P1 profile → 07_post_p1_profile.json"),
        ("08_generate_p1_doc.py",      "P1 repair report → deliverable/stage_p1_repair_report.docx"),
        ("09_p2_accuracy_efficacy.py", "P2 surveillance-domain checks → 09_p2_findings.json + 09_p2_examples.json"),
        ("10_generate_p2_doc.py",      "P2 report → deliverable/stage_p2_accuracy_efficacy.docx"),
        ("11_generate_methodology_journal.py", "This document → deliverable/methodology_journal.docx"),
        ("12_test_cleaned_p1.py",      "P1 invariant test suite (42 assertions) — fails if any P1 rule's claimed outcome doesn't hold on the file"),
        ("13_test_p2_findings.py",     "P2 reproducibility test suite (21 assertions) — fails if any P2 finding count drifts from the file"),
        ("14_p3_metrics.py",           "P3 surveillance DQ scorecard computation → 14_p3_metrics.json (44 metrics across 7 themes M1-M7)"),
        ("15_generate_p3_doc.py",      "P3 scorecard document → deliverable/stage_p3_metrics_scorecard.docx"),
        ("16_generate_exec_summary.py", "One-page Executive Summary cover → deliverable/00_executive_summary.docx"),
    ])
    p(doc, "Pinned versions: Python 3.14.0 · pandas 3.0.2 · numpy 2.4.4 · python-docx 1.2.0", size=9, italic=True)


def epilogue(doc):
    add_heading(doc, "12. If the Interviewer Asks Me to Pick the Three Most Important Things I Learned", level=1)
    bullet(doc, "The right unit of work is a NAMED rule, not a notebook cell. Naming the rule forces me to state the rationale, the action, and the count — which is exactly what an interviewer (or a regulator) needs to interrogate the work.")
    bullet(doc, "Big aggregate numbers are often parser artifacts, not data defects. The 'column-shift contamination' that looked like a major Stage 1 finding turned out to be a parser symptom. Fix the structural cause first; re-measure; only then describe what's left as data corruption.")
    bullet(doc, "Open questions are a finding, not a weakness. Listing three specific items the data producer needs to confirm (Account-CPF collapse, 5,142 mislabelled events, DD-MM date ambiguity) communicates more domain understanding than picking an interpretation and rolling with it.")


def main():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    cover(doc)
    opening_frame(doc)
    first_contact(doc)
    first_profile(doc)
    evidence_extraction(doc)
    power_query_moment(doc)
    stage_p1_thinking(doc)
    stage_p1_value_normalisation(doc)
    stage_p2_strategy(doc)
    stage_p2_execution(doc)
    decision_register(doc)
    reproducibility_map(doc)
    epilogue(doc)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED. Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
