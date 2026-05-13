"""
Stage P2 — Accuracy & Efficacy Report (Word doc).

Each finding is rendered with: surveillance rationale (the regulator's question),
detection logic, count, and 3–5 row-level examples (where applicable). Findings
are grouped by theme (A through G).
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).resolve().parent
DELIV = HERE.parent / "deliverable"
FINDINGS = json.loads((HERE / "09_p2_findings.json").read_text(encoding="utf-8"))
EXAMPLES = json.loads((HERE / "09_p2_examples.json").read_text(encoding="utf-8"))
POSTP1 = json.loads((HERE / "07_post_p1_profile.json").read_text(encoding="utf-8"))

OUT = DELIV / "stage_p2_accuracy_efficacy.docx"
OUT_FALLBACK = DELIV / "stage_p2_accuracy_efficacy_v2.docx"


# -------- helpers (self-contained copy of the docx primitives used in 03/08) --
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10)
    return p


def add_kv_table(doc, pairs, col0_width=2.0, col1_width=4.5):
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    t.columns[0].width = Inches(col0_width)
    t.columns[1].width = Inches(col1_width)
    for i, (k, v) in enumerate(pairs):
        c0, c1 = t.rows[i].cells[0], t.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(str(k))
        r0.bold = True
        r0.font.size = Pt(9)
        r1 = c1.paragraphs[0].add_run(str(v))
        r1.font.size = Pt(9)


def severity_box(doc, label, severity, text):
    color = {"P1-CRITICAL": "C0392B", "P1-HIGH": "E67E22",
             "P2-CRITICAL": "9B59B6", "P2-HIGH": "2980B9",
             "INFO": "7F8C8D", "RESOLVED": "27AE60"}.get(severity, "7F8C8D")
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(1.2)
    t.columns[1].width = Inches(5.3)
    cell_label, cell_text = t.rows[0].cells[0], t.rows[0].cells[1]
    cell_label.text = ""
    cell_text.text = ""
    set_cell_shading(cell_label, color)
    set_cell_shading(cell_text, "F4F6F8")
    rl = cell_label.paragraphs[0].add_run(severity)
    rl.bold = True
    rl.font.size = Pt(8)
    rl.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = cell_text.paragraphs[0]
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.size = Pt(9)


def add_example_table(doc, rows, fields=None, max_chars=28):
    if not rows:
        p = doc.add_paragraph()
        r = p.add_run("(no row-level examples for this check — either count was 0 or check is aggregate)")
        r.italic = True
        r.font.size = Pt(9)
        return
    if fields is None:
        fields = list(rows[0].keys())
    cols = ["Field"] + [f"Example {i+1}" for i in range(len(rows))]
    t = doc.add_table(rows=len(fields) + 1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(cols):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(c)
        r.bold = True
        r.font.size = Pt(8)
        set_cell_shading(cell, "1F3A5F")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, fname in enumerate(fields, start=1):
        cell = t.rows[i].cells[0]
        cell.text = ""
        r = cell.paragraphs[0].add_run(fname)
        r.bold = True
        r.font.size = Pt(8)
        for j, row in enumerate(rows, start=1):
            val = str(row.get(fname, ""))
            if len(val) > max_chars:
                val = val[:max_chars] + "…"
            cell = t.rows[i].cells[j]
            cell.text = ""
            rr = cell.paragraphs[0].add_run(val)
            rr.font.name = "Consolas"
            rr.font.size = Pt(8)


# -------- finding renderer ----------------------------------------------------
def render_finding(doc, key, severity_default="P2-HIGH"):
    f = FINDINGS.get(key)
    if not f:
        return
    sev = f.get("severity", severity_default)
    add_heading(doc, f"{key} — {f['name']}", level=3)
    add_kv_table(doc, [
        ("Theme", f["theme"]),
        ("Surveillance rationale", f["rationale"]),
        ("Detection logic", f["detection"]),
        ("Rows / hits", f"{f['count']:,}"),
    ])
    # extra context lines
    extras = {k: v for k, v in f.items() if k not in {"theme", "name", "rationale", "detection", "count", "severity"}}
    if extras:
        for k, v in extras.items():
            if isinstance(v, (int, float, str)):
                add_para(doc, f"{k}: {v}", italic=True, size=9)
            elif isinstance(v, dict):
                add_para(doc, f"{k}:", italic=True, size=9)
                for k2, v2 in v.items():
                    add_para(doc, f"  • {k2}: {v2}", size=9)
    severity_box(doc, f"Severity assessment", sev,
                 "Real surveillance impact — see rationale above.")
    examples = EXAMPLES.get(key) or []
    if examples:
        add_para(doc, "Row-level examples (first 5):", bold=True, size=9)
        # If examples are simple dicts, render
        add_example_table(doc, examples[:5])
    doc.add_paragraph()


# -------- sections ------------------------------------------------------------
def build_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TD Securities — Trade Surveillance Case Study")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Stage P2 — Accuracy & Efficacy Report")
    rs.italic = True
    rs.font.size = Pt(14)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs2 = sub2.add_run("Surveillance-domain checks on the P1-cleaned dataset")
    rs2.italic = True
    rs2.font.size = Pt(11)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared by: Shravan Challa\n").font.size = Pt(10)
    meta.add_run(f"Date: {date.today().isoformat()}\n").font.size = Pt(10)
    meta.add_run("Companions: Stage 1 Audit Trail, Stage P1 Repair Report\n").font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    pr = p.add_run("Purpose. ")
    pr.bold = True
    p.add_run(
        "Stage P1 fixed the mechanical defects that prevent ingestion. Stage P2 asks the more important question: "
        "even if the data is mechanically clean, can the surveillance system actually do its job with it? The brief "
        "describes this priority as 'open-ended' and 'based on your imagination' — so the checks below are framed "
        "as a regulator-shaped question for each: what surveillance rule does this data need to support, and does "
        "the data support it? Every finding is rendered with rationale (the WHY), detection logic (the HOW), a count, "
        "and row-level evidence where applicable."
    )
    doc.add_page_break()


def build_exec_summary(doc):
    add_heading(doc, "Executive Summary", level=1)

    add_para(doc, "Top-line read on the P1-cleaned dataset (242,429 rows):", bold=True)
    add_kv_table(doc, [
        ("Most consequential DQ finding", "Instrument↔ISIN mapping unstable: 222 of ~373 Instruments map to multiple ISINs (~60%). Every per-instrument surveillance rule is unreliable until this is resolved."),
        ("Most consequential attribution finding", "39,115 of 60,299 Fills (64.9%) carry NO Trader. Regulators require human attribution on executions; this is a Reg-best-execution / MAR-attribution gap."),
        ("Most consequential lifecycle finding", "1,748 parent orders show aggregate Fill volume > parent's TotalVolume. Textbook execution-integrity violation — every regulator exam tests for this."),
        ("Most consequential lineage finding", "22,287 child messages reference an upstream message that does not exist in the file (8,162 orphaned LinkMessageIds + 14,125 orphaned ParentOrderIds)."),
        ("Real surveillance positives flagged", "4,632 wash-trade signatures (same Account, opposite sides on same Instrument, within 5 seconds). These would be reviewed by surveillance analysts."),
        ("Baseline-noise concern", "149,798 rows (62%) show non-PROP Account self-crossing. This either reflects a synthetic-seeding choice or the Account column is a copy of CounterPartyFirm. Either way it floods wash-trade detection with false positives."),
    ])
    doc.add_page_break()


def build_methodology(doc):
    add_heading(doc, "1. Methodology", level=1)

    add_heading(doc, "1.1 Framing", level=2)
    add_para(doc,
             "P2 is the open-ended priority. The brief invited 'imagination' here. The lens I'm using is the surveillance lens: "
             "for each check, the framing is 'what surveillance rule does this support, and would a regulator be satisfied with the data?'. "
             "That's a higher bar than generic DQ — it forces every check to have a stated downstream consumer.", size=10)

    add_heading(doc, "1.2 Themes", level=2)
    add_kv_table(doc, [
        ("A. Referential integrity (lineage)", "Can surveillance reconstruct the order lifecycle from this file?"),
        ("B. Order lifecycle coherence", "Do the events form physically possible sequences (no fills before orders, no fills after cancels, no over-fills)?"),
        ("C. Instrument / ISIN coherence", "Is the symbol-to-ISIN mapping stable enough to support per-instrument rules?"),
        ("D. Account / Firm relationships", "When Account ≡ CounterPartyFirm, is it prop activity (legit) or agency self-cross (wash-trade candidate)?"),
        ("E. Trader attribution", "Can every execution be tied to a human (or attested algo)?"),
        ("F. Surveillance signatures", "Direct detections — cancel ratios, wash-trade signatures, etc."),
        ("G. Statistical anomalies", "Per-instrument price/volume outliers."),
    ])

    add_heading(doc, "1.3 Severity scheme", level=2)
    add_kv_table(doc, [
        ("P1-CRITICAL", "Blocks surveillance from running at all on affected rows."),
        ("P1-HIGH", "Surveillance runs, but specific rules silently fail on affected rows."),
        ("P2-CRITICAL", "Surveillance produces output, but the output is wrong (false negatives or false positives)."),
        ("P2-HIGH", "Surveillance produces output, but with high noise (more analyst review time)."),
        ("INFO", "Observation only — useful for producer confirmation but not a defect."),
    ])

    add_heading(doc, "1.4 Note on synthetic-data interpretation", level=2)
    add_para(doc,
             "The dataset is synthetic and seeded with intentional defects. Many P2 findings could be either (a) genuine "
             "modelling of real-world surveillance signals or (b) artefacts of the seeding program. The checks are surveillance-honest "
             "either way: they would fire on production data too. Where a finding strongly suggests a seeded shape (e.g., counts that "
             "exactly match another count), I note it.", size=10)
    doc.add_page_break()


def build_theme(doc, theme_key: str, theme_title: str, findings_keys: list[str], theme_summary: str):
    add_heading(doc, theme_title, level=1)
    add_para(doc, theme_summary, size=10)
    for k in findings_keys:
        render_finding(doc, k)
    doc.add_page_break()


def build_cross_cutting(doc):
    add_heading(doc, "9. Cross-Cutting Observations", level=1)

    add_heading(doc, "9.1 Three coincidences worth confirming with the producer", level=2)
    add_bullet(doc, "Count of New Orders with parent references (A4) = 5,142, which is the exact count of Replace Orders in the file. Could indicate: (a) seeded mislabelling of some Replace events as 'New Order', or (b) a pure coincidence. Cheap to confirm.")
    add_bullet(doc, "Account ≡ CounterPartyFirm on 97.4% of rows (236,043 of 242,429). The Stage-1 profile shows 1,043 distinct Account values vs 100 distinct CounterPartyFirm values, so Account IS more granular in principle — but in 97% of rows, the more-granular value has been collapsed back to the firm-level value. This pattern looks more like a column-copy defect upstream than a property of real prop trading flow.")
    add_bullet(doc, "65% of Fills have no Trader. In production this would mean two-thirds of executions are unattributable to a human, which is a Reg-best-execution / MAR finding. Almost certainly seeded — but the structure of the test is realistic.")

    add_heading(doc, "9.2 What surveillance rules cannot be confidently run on this data", level=2)
    add_bullet(doc, "Per-instrument price-deviation rules (G1) — symbol/ISIN mapping unstable (C1).")
    add_bullet(doc, "Self-trade / wash-trade detection (F2) — flooded by the 149,798-row baseline of non-PROP self-crosses (D1b).")
    add_bullet(doc, "Execution-attribution (best-ex / MAR) — fails on 65% of Fills (E1).")
    add_bullet(doc, "Lifecycle reconstruction for child events — fails for ~9% of references (A1+A2).")

    add_heading(doc, "9.3 What surveillance rules CAN run with high confidence on this data", level=2)
    add_bullet(doc, "Aggregate fill > order volume (B2) — clean numeric check, no dependency on mapping or attribution.")
    add_bullet(doc, "Time-order checks on child vs parent (B1, B3) — clean check given resolvable parent references.")
    add_bullet(doc, "PK-uniqueness / duplicate detection (already P1).")
    add_bullet(doc, "Volume outlier detection (G2) — small count of hits, low noise.")

    doc.add_page_break()


def build_open(doc):
    add_heading(doc, "10. Open Questions / Required Producer Input", level=1)
    add_bullet(doc, "Confirm whether the Account-≡-CounterPartyFirm collapse (97.4% of rows) is a seeded shape or a column-copy defect upstream.")
    add_bullet(doc, "Confirm whether the 5,142 'New Orders with parent references' are mislabelled Replace events.")
    add_bullet(doc, "Confirm whether the 222 instruments with multiple ISINs reflect a real corporate-action timeline (rare on a single trading day) or seeded mapping instability.")
    add_bullet(doc, "Confirm whether 'Trader is null' on Fills is intended to model algo executions in production (in which case the algo-ID should be a separate non-null attribution field).")
    add_bullet(doc, "Confirm preferred cancel-ratio threshold for the F1 spoofing signature (we used >5x as a placeholder; surveillance teams typically tune per desk).")
    add_bullet(doc, "Confirm preferred IQR multiplier and per-instrument minimum-sample-size threshold for the G1 price-outlier rule.")
    doc.add_page_break()


def build_recommendations(doc):
    add_heading(doc, "11. Recommendations for Stage P3 (Metrics)", level=1)
    add_bullet(doc, "Build the DQ scorecard around six metrics that map directly to surveillance capability: completeness, attribution, lineage, lifecycle integrity, mapping stability, signal-to-noise.")
    add_bullet(doc, "Express each as a percentage with a clear denominator. E.g., Attribution Rate = (Fills with non-null Trader) / (Fills). The team can then run the same script on each new vendor file and watch the trend.")
    add_bullet(doc, "Emit a 'surveillance readiness' top-line score that aggregates the six dimensions with explicit weights, so a non-technical reader can see at a glance whether the file is ingestable AND analysable.")
    add_bullet(doc, "Promote the 'Unmatched to Market Data from Vendor' Flag-column population (~6,623 rows) to a first-class metric — that flag is itself a surveillance gap and should be tracked across files.")
    doc.add_page_break()


def build_appendix(doc):
    add_heading(doc, "Appendix A — Reproducibility", level=1)
    add_para(doc, "Run these scripts in order from case-study/analysis/ to regenerate every artifact:", size=10)
    add_bullet(doc, "01_profile.py            — Stage 1 raw profile")
    add_bullet(doc, "02_extract_examples.py   — Stage 1 row-level evidence")
    add_bullet(doc, "04_line_count_audit.py   — line-ending diagnostic")
    add_bullet(doc, "05_repair_structural.py  — P1 structural repair")
    add_bullet(doc, "06_normalize_values.py   — P1 value normalisation")
    add_bullet(doc, "07_post_p1_profile.py    — post-P1 profile")
    add_bullet(doc, "09_p2_accuracy_efficacy.py — P2 surveillance-domain checks (this stage)")
    add_bullet(doc, "10_generate_p2_doc.py    — this document")


def main():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    build_cover(doc)
    build_exec_summary(doc)
    build_methodology(doc)

    build_theme(doc, "A", "2. Theme A — Referential Integrity (Lineage)",
                ["A1", "A2", "A3", "A4"],
                "Can surveillance reconstruct the order lifecycle by chasing LinkMessageId / ParentOrderId references back to a parent New Order? Where these references break, every chain-of-events surveillance rule fails silently for the affected rows.")

    build_theme(doc, "B", "3. Theme B — Order Lifecycle Coherence",
                ["B1", "B2", "B3", "B4"],
                "Do child events follow their parent in a physically-possible sequence? Fills cannot occur before the parent New Order. Cumulative Fill volume cannot exceed the parent's TotalVolume. Fills cannot land after the order is Cancelled. Each of these is a textbook regulator-exam check.")

    build_theme(doc, "C", "4. Theme C — Instrument / ISIN Coherence",
                ["C1", "C2"],
                "Every per-instrument surveillance rule (price-outlier, position-aggregation, market-data join) requires a stable Instrument↔ISIN mapping. If a single Instrument resolves to multiple ISINs on a single trading day, downstream rules cannot tell whether they're looking at one security or many.")

    build_theme(doc, "D", "5. Theme D — Account / Firm Relationships",
                ["D1", "D1b"],
                "Account-equals-CounterPartyFirm is a wash-trade signature on agency flow but expected on prop-trading accounts. The right read depends on the account-type prefix (PROP, INST, CLNT, BRKR, etc.).")

    build_theme(doc, "E", "6. Theme E — Trader Attribution",
                ["E1"],
                "Regulators require human attribution on executions. A Fill with no Trader is unattributable and breaks Reg-best-execution / MAR-attribution requirements.")

    build_theme(doc, "F", "7. Theme F — Surveillance Signatures",
                ["F1", "F2"],
                "Direct surveillance detections — cancel-to-new-order ratios per Trader (spoofing/layering signature) and wash-trade signatures (same Account, opposite sides on same Instrument within seconds).")

    build_theme(doc, "G", "8. Theme G — Statistical Anomalies",
                ["G1", "G2"],
                "Per-instrument outlier detection using a 3·IQR fence. Surveillance uses these as a low-cost first pass to triage rows for analyst review.")

    build_cross_cutting(doc)
    build_open(doc)
    build_recommendations(doc)
    build_appendix(doc)

    try:
        doc.save(OUT)
        print(f"Wrote: {OUT}")
    except PermissionError:
        doc.save(OUT_FALLBACK)
        print(f"PRIMARY LOCKED. Wrote fallback: {OUT_FALLBACK}")


if __name__ == "__main__":
    main()
